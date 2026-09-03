"""P2-27 (P26b-1, B112): the buyer-docx parser is body-first — headers and
footers are recorded as present (intake reads them; they bear no slots),
and a text box's words join the section they float in."""

from pathlib import Path

from engine.structure import parse_buyer_docx
from tests.fixtures.docx_twins import TEXT_BOX_DIRECTIVE

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"


def test_header_footer_presence_is_recorded():
    parsed = parse_buyer_docx(FIXTURES / "parts-twin.docx")
    assert any(w.startswith("header/footer text present (2 line(s)")
               for w in parsed.warnings)


def test_text_box_words_join_their_section(tmp_path):
    # Hand count on the parts twin: section 1 owns the question table
    # (a header slot + 2 open rows), section 2 is prose (1 ask) — 4 slots;
    # the text box floats in section 1, whose instructions a header slot
    # does not render. So the join is proven on a synthetic document
    # where the box is the ONLY prose under its numbered heading.
    parsed = parse_buyer_docx(FIXTURES / "parts-twin.docx")
    assert parsed.slot_count == 4
    two = next(s for s in parsed.slots if s.get("ref_id") == "2")
    assert two["question_text"] == "Describe your project governance model."

    from docx import Document
    from docx.oxml import parse_xml
    from tests.fixtures.docx_twins import _VML_TEXT_BOX, _table
    doc = Document()
    doc.add_heading("1. Section", level=1)
    _table(doc, [["Question", "Response"], ["Describe your approach.", ""]])
    doc.add_heading("2. Data Residency", level=1)
    anchor = doc.add_paragraph("")  # empty anchor: the box is the only words
    anchor._p.append(parse_xml(_VML_TEXT_BOX.format(text=TEXT_BOX_DIRECTIVE)))
    path = tmp_path / "boxed.docx"
    doc.save(path)
    parsed = parse_buyer_docx(path)
    two = next(s for s in parsed.slots if s.get("ref_id") == "2")
    assert two["question_text"] == TEXT_BOX_DIRECTIVE
