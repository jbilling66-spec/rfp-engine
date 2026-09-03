"""P2-27 (P26b-1, B112): the one implementation of "read the parts a
body walk never sees" — headers, footers, text boxes — proven on the
parts twin, whose planted strings are hand-known from
`tests/fixtures/docx_twins.py`."""

from pathlib import Path

from docx import Document

from engine.structure import header_footer_text, text_box_text
from tests.fixtures.docx_twins import (
    FOOTER_TEXT, HEADER_DIRECTIVE, TEXT_BOX_DIRECTIVE,
)

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"


def test_header_and_footer_are_read_in_order():
    document = Document(str(FIXTURES / "parts-twin.docx"))
    assert header_footer_text(document) == [
        ("header", HEADER_DIRECTIVE), ("footer", FOOTER_TEXT),
    ]


def test_text_box_under_its_anchor_paragraph():
    document = Document(str(FIXTURES / "parts-twin.docx"))
    anchors = [p for p in document.paragraphs
               if p.text.startswith("Answer each question")]
    assert len(anchors) == 1
    assert text_box_text(anchors[0]._p) == [TEXT_BOX_DIRECTIVE]
    assert text_box_text(document.paragraphs[0]._p) == []


def test_a_plain_document_has_no_parts():
    document = Document(str(FIXTURES / "qform-twin.docx"))
    assert header_footer_text(document) == []
    assert all(text_box_text(p._p) == [] for p in document.paragraphs)
