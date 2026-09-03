"""P2-30 (P26b-1, B112): a ragged buyer table — a body row with fewer
cells than the header — refuses with a typed StructureError naming the
table and row, at parse time AND at the write-back re-derivation
(`question_cell_map`), never a bare IndexError. The ragged row is made
by removing one <w:tc> from the saved OOXML, which is how real ragged
tables arise (editors other than Word)."""

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from engine.structure import parse_buyer_docx
from engine.structure.docx_buyer import question_cell_map
from engine.structure.parse import StructureError
from tests.fixtures.docx_twins import _table


def _ragged_doc(tmp_path: Path, *, question_table: bool) -> Path:
    doc = Document()
    doc.add_heading("1. Section", level=1)
    _table(doc, [["Question", "Response"], ["Describe your approach.", ""]])
    rows = ([["Question", "Response"], ["Describe your team.", ""], ["Describe your tooling.", ""]]
            if question_table else
            [["Role", "Name", "% Allocation"], ["PM", "", ""], ["Architect", "", ""]])
    _table(doc, rows)
    table = doc.tables[1]
    tr = table.rows[2]._tr
    tr.remove(tr.findall(qn("w:tc"))[-1])  # row 2 loses its last cell
    path = tmp_path / "ragged.docx"
    doc.save(path)
    return path


def test_ragged_question_table_refuses_typed_at_parse(tmp_path):
    path = _ragged_doc(tmp_path, question_table=True)
    with pytest.raises(StructureError) as info:
        parse_buyer_docx(path)
    assert "table 1 row 2: 1 cells, header has 2" in str(info.value)
    assert "ragged" in str(info.value)


def test_ragged_fill_in_table_refuses_typed_at_parse(tmp_path):
    path = _ragged_doc(tmp_path, question_table=False)
    with pytest.raises(StructureError) as info:
        parse_buyer_docx(path)
    assert "table 1 row 2: 2 cells, header has 3" in str(info.value)


def test_ragged_table_refuses_typed_at_write_back_rederivation(tmp_path):
    path = _ragged_doc(tmp_path, question_table=True)
    with pytest.raises(StructureError) as info:
        question_cell_map(path)
    assert "table 1 row 2" in str(info.value)
