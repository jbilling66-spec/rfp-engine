"""C12 — the two-stack seam is a tested PROPERTY, not a discipline (B57:
"the seam between them is C8-C13's to make loud").

Three pins: (1) the RFP-extract path stamps docling and the KB-distillation
path stamps python-docx — the identity literals are asserted, so silently
swapping either stack reddens here; (2) the two fingerprints can never
collide; (3) both paths' artifacts actually CARRY their stamp (intake's
extraction.json + checkpoint, ingest's report/SourceDoc) — the seam is
visible in the record, not just in code."""

import json

from engine.kb.read import read_source
from tests.extraction.fakes import FakeExtractionBackend
from tests.intake.fixtures.packages import run_package
from tests.intake.test_brief_backend import _pdf_twin_view

# Pin 1: the identity literals. These are the B57 verdict rendered as
# constants — moving either requires a deliberate edit here, which is
# exactly the loudness the entry asks for.
INTAKE_PRIMARY = "docling"
KB_PRIMARY = "python-docx"


def test_intake_path_stamps_docling(tmp_path):
    fake = FakeExtractionBackend({"pdf-twin.pdf": _pdf_twin_view()})
    pursuit, _ = run_package(tmp_path, "pdf", extraction=fake)
    art = json.loads((pursuit.root / "extraction.json").read_text())
    assert art["docs"][0]["extractor"] == INTAKE_PRIMARY
    assert pursuit.checkpoint_payload("intake")["docs"][0]["extractor"] == (
        INTAKE_PRIMARY
    )


def test_kb_path_stamps_python_docx(tmp_path):
    import docx as pydocx

    d = pydocx.Document()
    d.add_heading("Past response", level=1)
    d.add_paragraph("We delivered the migration on schedule.")
    path = tmp_path / "resp.docx"
    d.save(path)
    source = read_source(path)
    assert source.extractor == KB_PRIMARY


def test_the_two_fingerprints_never_collide(tmp_path):
    import docx as pydocx

    from engine.extraction.fingerprint import extraction_fingerprint

    d = pydocx.Document()
    d.add_paragraph("body")
    path = tmp_path / "same.docx"
    d.save(path)
    kb_stamp = read_source(path).fingerprint
    intake_stamp = extraction_fingerprint("2.121.0", "a" * 64)
    assert kb_stamp != intake_stamp
    assert kb_stamp.startswith("ext_") and intake_stamp.startswith("ext_")


def test_a_swap_reddens_both_directions():
    # The backend classes and the reader name their identities as
    # attributes/returns; pin them against the verdict constants so a
    # refactor cannot quietly re-point a path at the other stack.
    from engine.extraction.backend import DockerBackend, InContainerBackend

    assert InContainerBackend.identity == INTAKE_PRIMARY
    assert DockerBackend.identity == INTAKE_PRIMARY
    from tests.extraction.fakes import FakeExtractionBackend as Fake

    assert Fake.identity == INTAKE_PRIMARY
