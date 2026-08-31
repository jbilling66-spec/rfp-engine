"""Corpus trait proofs (C4, B51) — offline, raw-library read-back.

Each built document is read back with the library that understands it
(python-docx / openpyxl-free zip / pypdf) and compared against the
committed answer key evals/extraction-gate/ground_truth.json. The key
and the builders are duplicated on purpose: drift on either side
reddens here. The scanned twin's trait test rides the container-only
roster (PIL); everything else runs everywhere.
"""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pytest

from engine.extraction import corpus

ROOT = Path(__file__).resolve().parents[2]
GROUND_TRUTH = json.loads(
    (ROOT / "evals" / "extraction-gate" / "ground_truth.json").read_text()
)


@pytest.fixture(scope="module")
def built(tmp_path_factory):
    return corpus.build_corpus(tmp_path_factory.mktemp("gate-corpus"))


def _docx_tables(path: Path):
    import docx

    return docx.Document(str(path)).tables


def _grid(table):
    return [[cell.text for cell in row.cells] for row in table.rows]


def _fills(table):
    from docx.oxml.ns import qn

    found = {}
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            for shd in cell._tc.iter(qn("w:shd")):
                found.setdefault(shd.get(qn("w:fill")), set()).add((r, c))
    return found


def test_complex_table_grid_matches_the_answer_key(built):
    truth = GROUND_TRUTH["complex-tables-twin.docx"]["tables"]
    ct1, ct2 = _docx_tables(built["complex-tables-twin.docx"])
    assert _grid(ct1) == truth["CT-1"]["grid"]
    assert _grid(ct2) == truth["CT-2"]["grid"]


def test_complex_table_merges_are_real_merges(built):
    """Merged positions are the same underlying cell, not lookalike text."""
    ct1 = _docx_tables(built["complex-tables-twin.docx"])[0]
    for (ar, ac), (br, bc) in GROUND_TRUTH["complex-tables-twin.docx"]["tables"]["CT-1"]["merges"]:
        assert ct1.rows[ar].cells[ac]._tc is ct1.rows[br].cells[bc]._tc, (
            f"({ar},{ac}) and ({br},{bc}) must be one merged cell"
        )


def test_shading_is_present_where_the_key_says(built):
    for doc_name in ("complex-tables-twin.docx", "response-form-twin.docx"):
        for table_id, truth in GROUND_TRUTH[doc_name]["tables"].items():
            if not truth["fills"]:
                continue
            table = _docx_tables(built[doc_name])[0]
            found = _fills(table)
            for fill, positions in truth["fills"].items():
                assert fill in found, f"{doc_name}/{table_id}: fill {fill} absent"
                for r, c in positions:
                    assert (r, c) in found[fill], (
                        f"{doc_name}/{table_id}: fill {fill} missing at ({r},{c})"
                    )


def test_form_grid_merge_and_comment_anchor(built):
    truth = GROUND_TRUTH["response-form-twin.docx"]
    form = _docx_tables(built["response-form-twin.docx"])[0]
    assert _grid(form) == truth["tables"]["F-1"]["grid"]
    (a, b), = truth["tables"]["F-1"]["merges"]
    assert form.rows[a[0]].cells[a[1]]._tc is form.rows[b[0]].cells[b[1]]._tc
    # comment presence asserted at the zip level — API-version independent
    with zipfile.ZipFile(built["response-form-twin.docx"]) as zf:
        assert "word/comments.xml" in zf.namelist()
        comments = zf.read("word/comments.xml").decode("utf-8")
    assert truth["comment_anchor"]["text"] in comments


def test_control_hierarchy_matches_the_answer_key(built):
    import docx

    doc = docx.Document(str(built["firm-control-twin.docx"]))
    headings = [
        [int(p.style.name.split()[-1]), p.text]
        for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]
    assert headings == GROUND_TRUTH["firm-control-twin.docx"]["hierarchy"]


def test_multicolumn_pdf_carries_both_columns(built):
    from pypdf import PdfReader

    text = PdfReader(built["multicolumn-twin.pdf"]).pages[0].extract_text()
    truth = GROUND_TRUTH["multicolumn-twin.pdf"]
    assert truth["title"].split(" - ")[0] in text
    for line in truth["reading_order"]:
        assert line.split()[0] in text  # every RO-nn marker present


def test_table_pdf_carries_every_cell_text(built):
    from pypdf import PdfReader

    text = PdfReader(built["table-pdf-twin.pdf"]).pages[0].extract_text()
    truth = GROUND_TRUTH["table-pdf-twin.pdf"]["grid"]
    for cell in {c for row in truth for c in row}:
        assert cell in text, f"cell text {cell!r} missing from the ruled PDF"


def test_500_page_pdf_has_500_pages(built):
    from pypdf import PdfReader

    assert len(PdfReader(built["500-page-twin.pdf"]).pages) == GROUND_TRUTH[
        "500-page-twin.pdf"
    ]["pages"]


def test_encrypted_pdf_is_really_encrypted(built):
    from pypdf import PdfReader

    assert PdfReader(built["encrypted-twin.pdf"]).is_encrypted


def test_truncated_pdf_is_really_broken(built):
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    with pytest.raises(PyPdfError):
        PdfReader(built["truncated-twin.pdf"]).pages[0].extract_text()


def test_media_docs_embed_a_generated_png(built):
    for name in ("logo-twin.docx", "signature-twin.docx"):
        with zipfile.ZipFile(built[name]) as zf:
            images = [n for n in zf.namelist() if n.startswith("word/media/")]
            assert images, f"{name}: no embedded media part"
            assert zf.read(images[0])[:8] == b"\x89PNG\r\n\x1a\n"


def test_builders_are_deterministic(built, tmp_path):
    """Byte-equal on rebuild — except the encrypted twin, whose cipher
    salts are random by design (excluded, documented here)."""
    rebuilt = corpus.build_corpus(tmp_path / "again")
    for name, path in built.items():
        if name == "encrypted-twin.pdf":
            continue
        assert rebuilt[name].read_bytes() == path.read_bytes(), f"{name} not deterministic"


def test_corpus_carries_no_real_world_tokens(built):
    """The tripwire scans committed files; this corpus is generated, so
    sweep it here through the same token list (count and path only)."""
    from tests.tripwire.tokens import scan_tokens

    tokens = scan_tokens()
    for name, path in built.items():
        data = path.read_bytes().lower()
        for token in tokens:
            assert token.lower().encode() not in data, f"real-world token in {name}"
