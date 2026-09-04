"""The firm_default lane over HTTP (P17/C10; P26a item 1, P1-27): the
writeback dispatcher reads the container's OWN source_mode, preview →
confirm fills the template, and the downloads doors tell the truth
about what remains — a partly filled response lists as a WORKING copy
under internal (its authoring scaffolding already stripped) and the
buyer download refuses naming what is owed; once every section is
drafted and the hand-completion record is complete, the buyer copy
lists under the buyer heading and downloads clean. The generated render
still REFUSES firm_default with the pointer (B75§1d)."""

import hashlib
import io
import json

import pytest
from docx import Document
from fastapi.testclient import TestClient

from engine.planning.plan import REFERENCE_DEFAULT
from engine.structure import merge_parsed, parse_default_template
from engine.web.server import create_app
from engine.workspace import PursuitDir
from tests.web.conftest import FIXED_AT, sign_in
from tests.helpers import plant_annotated, plant_freeze

PARA = "A synthetic executive summary paragraph."
PROSE_SECTIONS = [1, 2, 3, 4, 5, 6, 7, 8, 9, 12, 13, 14]
HAND = {
    "s-front-meta": {
        "prepared_for_client": "Synthetic Buyer Co",
        "rfp_title": "Synthetic RFP", "rfp_solicitation_number": "RFP-0001",
        "submitted_by": "The Firm", "date_of_submission": "2026-09-30",
        "primary_contact": "Pat Lead", "due_date_method": "portal"},
    "s-h11": [{"milestone": "Kickoff", "fee": "$1,000",
               "duration_weeks": "2"}],
    "s-h10": [{"client": "A synthetic utility", "scope": "Finance",
               "outcome": "Live on schedule"}],
    "s-h12-1": "Net 30 from invoice",
}


def _plant(ws, pursuit_id: str, *, all_prose: bool):
    pursuit = PursuitDir(ws, pursuit_id)
    parsed = parse_default_template(REFERENCE_DEFAULT)
    container = {"pursuit_id": pursuit_id, **merge_parsed([parsed])}
    pursuit.write_artifact("target_slots", container, name="slots.json")
    pursuit.checkpoint("path_b_outline", {
        "reference_sha256": hashlib.sha256(
            REFERENCE_DEFAULT.read_bytes()).hexdigest()})
    ids = [f"s-h{n:02d}" for n in PROSE_SECTIONS] if all_prose else ["s-h02"]
    plant_freeze(pursuit, "pursuit_plan", {
        "pursuit_id": pursuit_id, "path": "B_free_flow",
        "slots_ref": "slots.json", "status": "approved",
        "sections": [{"section_id": f"sec-{s}",
                      "slot_ids": [f"{s}-hdr", s]} for s in ids]})
    (pursuit.root / "drafts" / "draft.json").write_text(json.dumps({
        "plan_sha256": pursuit.file_sha256("plan.frozen.json"),
        "revision_n": 0,
        "sections": [{"section_id": f"sec-{s}", "status": "drafted",
                      "prose": PARA if s == "s-h02"
                      else f"Synthetic prose for {s}."} for s in ids]}))
    plant_annotated(pursuit)
    return pursuit


@pytest.fixture(scope="module")
def fill_client(tmp_path_factory):
    ws = tmp_path_factory.mktemp("web-fill") / "ws"
    app = create_app(ws, now=lambda: FIXED_AT)
    client = TestClient(app, base_url="http://127.0.0.1")
    client.__enter__()
    sign_in(client, "Fiona Filler")
    client.post("/api/pursuits", json={"pursuit_id": "pur_fill"})
    _plant(ws, "pur_fill", all_prose=False)
    client.post("/api/pursuits", json={"pursuit_id": "pur_full"})
    _plant(ws, "pur_full", all_prose=True)
    yield client, ws
    client.__exit__(None, None, None)


def _docx_texts(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_preview_confirm_fill_and_download(fill_client):
    """One section drafted, nothing hand-entered: the fill produces the
    working copy, WITHHOLDS the buyer copy, and both download doors say
    so."""
    client, ws = fill_client
    body = client.get(
        "/api/pursuits/pur_fill/writeback/preview").json()
    preview = body["files"][0]  # the uniform shape (P18/C6, B77§2 D4)
    decisions = {r["slot_id"]: r["decision"]
                 for r in preview["sections"]}
    assert decisions["s-h02"] == "filled"
    assert preview["confirmed_by"] == "(unconfirmed)"
    assert preview["buyer_copy_produced"] is False

    facts = client.post("/api/pursuits/pur_fill/writeback/confirm",
                        json={})
    assert facts.status_code == 200, facts.text
    confirmed = facts.json()
    out = confirmed["files"][0]
    assert out["confirmed_by"] == "Fiona Filler"
    assert out["working_copy"] == "exports/review/response-working.docx"
    assert len(out["remaining_by_hand"]) == 4
    # the firm_default bundle is ONE entry, and it is REFUSED — the
    # recorded reason names what remains and where the working copy is
    deliverables = confirmed["bundle"]["deliverables"]
    assert [d["lane"] for d in deliverables] == ["template_fill"]
    assert deliverables[0]["status"] == "refused"
    assert deliverables[0]["reason"].startswith("buyer copy withheld")
    assert "working copy at exports/review/response-working.docx" \
        in deliverables[0]["reason"]
    assert not (ws / "pur_fill" / "exports" / "submission" /
                "response.docx").exists()
    assert (ws / "pur_fill" / "exports" / "review" /
            "response-working.docx").exists()

    downloads = client.get("/api/pursuits/pur_fill/downloads").json()
    assert downloads["to_the_buyer"] == []
    assert "response-working.docx" in downloads["internal_do_not_send"]
    refused = client.get("/api/pursuits/pur_fill/download/response.docx")
    assert refused.status_code == 409
    assert "buyer copy withheld" in refused.json()["detail"]
    working = client.get(
        "/api/pursuits/pur_fill/download/response-working.docx")
    assert working.status_code == 200
    text = _docx_texts(working.content)
    assert "How to Use" not in text and "Firm Response Template" not in text
    assert PARA in text and "▸" in text  # guidance kept for the hand-drafter


def test_complete_fill_lists_and_downloads_the_buyer_copy(fill_client):
    """Every prose section drafted + the hand-completion record entered
    through the door: the buyer copy is produced, listed under the
    buyer heading, and downloads with no scaffolding in it."""
    client, ws = fill_client
    r = client.put("/api/pursuits/pur_full/writeback/hand-fill",
                   json={"values": HAND})
    assert r.status_code == 200, r.text
    assert all(s["status"] == "filled" for s in r.json()["slots"])
    facts = client.post("/api/pursuits/pur_full/writeback/confirm",
                        json={})
    assert facts.status_code == 200, facts.text
    confirmed = facts.json()
    assert confirmed["files"][0]["buyer_copy_produced"] is True
    entry = confirmed["bundle"]["deliverables"][0]
    assert entry["status"] == "produced"
    assert entry["facts_path"] == "exports/template-fill-facts.json"
    downloads = client.get("/api/pursuits/pur_full/downloads").json()
    assert downloads["to_the_buyer"] == ["response.docx"]
    got = client.get("/api/pursuits/pur_full/download/response.docx")
    assert got.status_code == 200
    text = _docx_texts(got.content)
    for marker in ("▸", "[", "How to Use", "Firm Response Template",
                   "Replace with", "Field"):
        assert marker not in text, marker
    assert "Synthetic Buyer Co" in text and "Kickoff" in text
    assert "Payment schedule & terms: Net 30 from invoice" in text


def test_generated_render_refuses_firm_default_with_the_pointer(
        fill_client):
    client, _ws = fill_client
    resp = client.post("/api/pursuits/pur_fill/export",
                       json={"lane": "submission"})
    assert resp.status_code == 409
    assert "FILLED template" in resp.json()["detail"]


def test_downloads_and_detail_name_what_is_withheld(fill_client):
    """P27 wave 1: the downloads model carries every refused deliverable
    with the record's own reason (the finish panel renders it without a
    409 round-trip), and the detail model names the finishing
    preconditions the panel's buttons key on — reviewable, the bundle
    summary, the hand-completion lane."""
    client, ws = fill_client
    r = client.post("/api/pursuits/pur_fill/writeback/confirm", json={})
    assert r.status_code == 200, r.text
    downloads = client.get("/api/pursuits/pur_fill/downloads").json()
    assert [d["name"] for d in downloads["refused"]] == ["response.docx"]
    assert downloads["refused"][0]["reason"].startswith("buyer copy withheld")
    detail = client.get("/api/pursuits/pur_fill").json()
    finishing = detail["finishing"]
    assert finishing["reviewable"] is True
    assert finishing["hand_fill_lane"] is True
    assert finishing["bundle"]["produced"] == 0
    assert finishing["bundle"]["refused"] == 1
    assert finishing["bundle"]["composed_by"] == "Fiona Filler"
    # a bare pursuit: nothing to finish, nothing composed, no lane
    client.post("/api/pursuits", json={"pursuit_id": "pur_bare"})
    bare = client.get("/api/pursuits/pur_bare").json()
    assert bare["finishing"] == {"reviewable": False, "bundle": None,
                                 "hand_fill_lane": False}
    assert client.get("/api/pursuits/pur_bare/downloads").json() == {
        "to_the_buyer": [], "internal_do_not_send": [], "refused": []}


# --------------------------------------------- P26c: the case block learns

def test_confirm_proposes_the_case_block_not_the_grid(tmp_path, monkeypatch):
    """P1-44: at writeback confirm the hand-typed case block becomes a
    corpus case-study proposal in the steward inbox — in the human's
    words, naming its slot and artifact; the metadata record, the
    pricing grid and the inline line are skipped and named; nothing
    enters the store until a steward accepts; a learner failure never
    refuses the confirm."""
    from engine.flywheel.proposals import ProposalStore
    from engine.kb.curation import merge_batch
    from engine.kb.store import KBStore

    ws = tmp_path / "ws"
    store = KBStore(ws / "kb")  # BEFORE create_app: the app's own KB
    app = create_app(ws, now=lambda: FIXED_AT)
    with TestClient(app, base_url="http://127.0.0.1") as client:
        sign_in(client, "Fiona Filler")
        client.post("/api/pursuits", json={"pursuit_id": "pur_case"})
        _plant(ws, "pur_case", all_prose=True)
        put = client.put("/api/pursuits/pur_case/writeback/hand-fill",
                         json={"values": HAND})
        assert put.status_code == 200, put.text
        confirmed = client.post("/api/pursuits/pur_case/writeback/confirm",
                                json={})
        assert confirmed.status_code == 200, confirmed.text
        flywheel = confirmed.json()["flywheel"]
        assert len(flywheel["proposals"]) == 1
        assert set(flywheel["skipped"]) == {"s-front-meta", "s-h11", "s-h12-1"}
        assert "P3-1" in flywheel["skipped"]["s-h11"]
        proposal = ProposalStore(store.root).read(flywheel["proposals"][0])
        assert proposal["source"] == {
            "door": "flywheel", "pursuit_id": "pur_case", "slot_id": "s-h10",
            "artifact": "exports/hand-fill.json", "operator": "Fiona Filler"}
        assert proposal["kind"] == "new_card" and proposal["target"] == "corpus"
        assert proposal["diff"]["doc_kind"]["after"] == "case_study"
        assert "A synthetic utility" in proposal["diff"]["body"]["after"]
        assert "$1,000" not in proposal["diff"]["body"]["after"]
        assert store.list_cards() == [], "the inbox, never the corpus"
        # the steward accepts: a corpus case study, human-authored
        merge_batch(store, [proposal["proposal_id"]], operator="Sam",
                    at=FIXED_AT)
        minted = store.list_cards()
        assert len(minted) == 1
        assert minted[0]["doc_kind"] == "case_study"
        assert minted[0]["content_origin"] == "human_authored"
        # a second confirm re-proposes nothing (content-derived ids)
        again = client.post("/api/pursuits/pur_case/writeback/confirm",
                            json={}).json()["flywheel"]
        assert again["proposals"] == flywheel["proposals"]
        assert len(ProposalStore(store.root).list()) == 1
        # a learner failure is reported, never a refusal of the bundle
        import engine.web.learn as learn_mod

        def boom(*_a, **_k):
            raise RuntimeError("hand-fill learner on fire")

        monkeypatch.setattr(learn_mod, "learn_from_writeback", boom)
        r = client.post("/api/pursuits/pur_case/writeback/confirm", json={})
        assert r.status_code == 200, r.text
        assert r.json()["flywheel"] == {
            "error": "RuntimeError: hand-fill learner on fire"}
        assert r.json()["bundle"]
