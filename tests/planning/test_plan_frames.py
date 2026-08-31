"""The firm reference: loader contract, frame shape, and digest folding.
The architect-script frame-drift assertions live here too once C9's
regexes exist — single-source with the derive-wire fixture."""

import pytest

from engine.llm import effective_config
from engine.llm.frames import wrap_reference
from engine.planning.plan import REFERENCE_DEFAULT
from engine.planning.reference import (
    ReferenceError,
    load_reference,
    render_reference,
)
from tests.planning.fixtures.plans import planning_extras


def test_committed_reference_loads_fourteen_sections():
    outline = load_reference(REFERENCE_DEFAULT)
    assert len(outline.sections) == 14
    ids = outline.ids()
    assert ids[0] == "cover-letter" and ids[1] == "executive-summary"
    assert ids[-1] == "appendices"
    assert len(ids) == len(set(ids))
    for section in outline.sections:
        assert section["purpose"]  # every section carries its one-liner


def test_render_reference_is_the_based_on_vocabulary():
    outline = load_reference(REFERENCE_DEFAULT)
    rendered = render_reference(outline)
    for section in outline.sections:
        assert f"- {section['id']} | {section['title']} | " in rendered


def test_reference_sections_carry_the_template_slots():
    """P16/C7: the reference is the PARSED firm template — each section
    names the TargetSlot ids an adapting outline section inherits."""
    outline = load_reference(REFERENCE_DEFAULT)
    assert outline.parsed.source_mode == "firm_default"
    for section in outline.sections:
        assert section["slot_ids"], section["id"]
        assert section["slot_ids"][0].endswith("-hdr")
    exec_summary = next(s for s in outline.sections
                        if s["id"] == "executive-summary")
    assert exec_summary["purpose"].startswith("the firm's standard")
    assert "▸" not in exec_summary["purpose"]  # guidance marker stripped


def test_malformed_reference_raises(tmp_path):
    """Template flavor (P16/C7): the markdown lane retired — a .md path
    refuses by suffix, a missing or headerless template refuses via the
    parser, all as ReferenceError (loader-is-the-contract)."""
    with pytest.raises(ReferenceError, match="markdown reference retired"):
        load_reference(tmp_path / "old-reference.md")

    with pytest.raises(ReferenceError, match="no template"):
        load_reference(tmp_path / "missing.docx")

    from docx import Document
    doc = Document()
    doc.add_paragraph("No numbered headings, no placeholders.")
    doc.add_paragraph("Payment terms: [ … ]")  # front matter only
    empty = tmp_path / "not-a-template.docx"
    doc.save(empty)
    with pytest.raises(ReferenceError):
        load_reference(empty)


def test_wrap_reference_frame_shape():
    framed = wrap_reference("- cover-letter | Cover Letter | One page.")
    assert framed.startswith('<firm_reference label="firm">\n')
    assert framed.endswith("\n</firm_reference>")


def test_frames_match_harness_regexes():
    """Single-source tripwire (strategies.py precedent): if a frame's
    shape changes, the derive-wire script and this assertion fail
    together, never silently apart."""
    from engine.llm.frames import wrap_brief_context, wrap_lead_context
    from tests.planning.fixtures.plans import BRIEF_RX, LEAD_RX, REFERENCE_RX

    assert REFERENCE_RX.search(wrap_reference("- a | A | p"))
    assert BRIEF_RX.search(wrap_brief_context("Buyer: X"))
    assert LEAD_RX.search(wrap_lead_context("feedback text"))


def test_reference_and_manifest_are_digest_visible():
    """B16 closed: both planning inputs are run variables — editing
    either changes config_digest, so two runs are only comparable when
    their planning inputs match."""
    from engine.runlog import config_digest

    base = config_digest(effective_config(extra=planning_extras()))
    tweaked = effective_config(extra=planning_extras())
    tweaked["planning"]["reference_sha256"] = "0" * 64
    assert config_digest(tweaked) != base
    assert "reference_sha256" in effective_config(
        extra=planning_extras()
    )["planning"]
