"""The docx twins hold their contract (P16/C1): byte-deterministic
builds, committed goldens byte-matched, trap properties actually
present (a drifted twin must fail HERE, before any parser golden can
silently stop exercising its lane), and extracted text tripwire-clean.
"""

import hashlib
from pathlib import Path

import pytest
from docx import Document

from tests.fixtures.docx_twins import GOLDENS, rebuild_all

FIXTURES_DIR = Path(__file__).resolve().parent


def _sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_docx_builds_are_byte_deterministic(tmp_path):
    rebuild_all(tmp_path)
    first = {name: _sha(tmp_path / name) for name in GOLDENS}
    rebuild_all(tmp_path)
    assert {name: _sha(tmp_path / name) for name in GOLDENS} == first


@pytest.mark.parametrize("name", sorted(GOLDENS))
def test_docx_rebuild_matches_committed_golden(tmp_path, name):
    GOLDENS[name](tmp_path / name)
    assert _sha(tmp_path / name) == _sha(FIXTURES_DIR / name), (
        f"{name}: builder and committed golden diverged"
    )


def test_docx_twin_text_is_tripwire_clean():
    """Binary goldens dodge the text tripwire (suffix scan) — extract
    their text and sweep the real-client token list directly. New
    goldens ride this parametrization automatically via GOLDENS."""
    from engine.intake.extract import extract
    from tests.tripwire.tokens import scan_tokens

    tokens = scan_tokens()
    for name in sorted(GOLDENS):
        text = extract(FIXTURES_DIR / name).text.lower()
        for token in tokens:
            assert token not in text, f"{name} carries real-world token {token!r}"


def test_config_template_is_pinned_to_the_builder_and_swept():
    """P16/C7: config/templates/firm-default-template.docx is the RUNTIME
    firm template (the adopter-config surface) and must stay
    byte-identical to the builder — one source of truth, no drift
    between what tests prove and what Path B parses. Its extracted text
    is swept here because the tripwire's binary-closure names it."""
    from engine.intake.extract import extract
    from tests.tripwire.tokens import scan_tokens

    config_path = (FIXTURES_DIR.parents[1]
                   / "config" / "templates" / "firm-default-template.docx")
    assert _sha(config_path) == _sha(FIXTURES_DIR / "template-twin.docx")
    text = extract(config_path).text.lower()
    for token in scan_tokens():
        assert token not in text


def _texts(doc):
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    return paragraphs, cells


def test_template_twin_carries_its_conventions():
    doc = Document(str(FIXTURES_DIR / "template-twin.docx"))
    paragraphs, cells = _texts(doc)
    h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    numbered = [t for t in h1 if t[:1].isdigit()]
    assert len(numbered) == 14                                   # the structure
    assert all("  " in t.split(".")[1][:3] or t.split(".  ")[0].isdigit()
               for t in numbered)                                # "N.  Title"
    assert h1[0] == "How to Use This Template"                   # front matter
    assert doc.tables[0].rows[0].cells[0].text == "Field"        # metadata table
    assert sum(1 for c in cells if c.startswith("▸ WHAT TO INCLUDE")) == 14
    assert any("Keep it to one page" in c for c in cells)        # brevity carrier
    assert any("Client: [" in c and "Outcome: [" in c for c in cells)  # case block
    header = next(t for t in doc.tables
                  if t.rows[0].cells[0].text == "Milestone")
    assert [c.text for c in header.rows[0].cells] == [
        "Milestone", "Fee", "Duration (weeks)"]                  # typed grid
    assert any(p.startswith("Payment schedule & terms:") for p in paragraphs)
    assert any(p.text == "Timeline & Milestones"
               and p.style.name == "Heading 2" for p in doc.paragraphs)


def test_outline_twin_carries_its_conventions():
    doc = Document(str(FIXTURES_DIR / "outline-twin.docx"))
    paragraphs, _ = _texts(doc)
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert h2 == ["1. Executive Summary", "2. Implementation Approach",
                  "3. Staffing Plan", "4. Change Management",
                  "5. Value-Added Services (Optional)"]
    assert any(p.text == "2.1 Project Timeline"
               and p.style.name == "Heading 3" for p in doc.paragraphs)
    joined = "\n".join(paragraphs)
    assert "shall not exceed two (2) pages" in joined            # page limit
    assert "thirty percent (30%)" in joined                      # eval weight
    assert "optional and may be omitted" in joined               # required:false
    assert "follow the section numbering" in joined              # mirror mandate


def test_qform_twin_carries_its_conventions():
    doc = Document(str(FIXTURES_DIR / "qform-twin.docx"))
    q_tables = [t for t in doc.tables
                if t.rows[0].cells[0].text == "Question"]
    assert len(q_tables) == 2
    empties = [r for t in q_tables for r in t.rows[1:]
               if r.cells[1].text == ""]
    assert len(empties) == 5                                     # open answers
    assert any("Do you subcontract" in r.cells[0].text
               for t in q_tables for r in t.rows)                # boolean opener
    assert any("How many" in r.cells[0].text
               for t in q_tables for r in t.rows)                # numeric opener
    assert any(r.cells[1].text.startswith("EXAMPLE:")
               for t in q_tables for r in t.rows)                # pre-filled trap
    refs = next(t for t in doc.tables
                if t.rows[0].cells[0].text == "Reference")
    assert len(refs.columns) == 3                                # record grid


def test_narrative_twin_carries_its_conventions():
    doc = Document(str(FIXTURES_DIR / "narrative-twin.docx"))
    filled = next(t for t in doc.tables
                  if t.rows[0].cells[0].text == "Phase")
    assert all(c.text for r in filled.rows for c in r.cells)     # fully filled —
    fillin = next(t for t in doc.tables                          # never a slot
                  if t.rows[0].cells[0].text == "Role")
    assert [r.cells[0].text for r in fillin.rows[1:]] == [
        "Project Manager", "Solution Architect", "Change Lead"]
    assert all(r.cells[2].text == "" for r in fillin.rows[1:])   # F3: empty column
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert h2 == ["1. Technical Approach", "2. Support Model"]   # embedded outline
