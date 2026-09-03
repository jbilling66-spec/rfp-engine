"""Byte-deterministic synthetic DOCX twins (P16/C1).

Fresh implementations for the same synthetic buyer as the xlsx twins
("Northwind Regional Health") carrying the response-vehicle shapes the
P16 parsers must handle — each convention observed in the real pen
corpus (read-to-design, B71/B73: shapes carried, text synthetic):

  template-twin.docx   the firm's 14-section default template (v1's
                       default-template structure, the owner's kickoff call B73§2):
                       numbered "N.  Title" H1s, "▸ WHAT TO INCLUDE"
                       guidance tables, "[ … ]" placeholder tables, a
                       front-matter metadata table, a one-page brevity
                       carrier, a case-study inner-field block, a typed
                       pricing grid, an inline bracketed body paragraph.
  outline-twin.docx    a buyer Word outline: mandated numbered sections
                       as headings with instruction paragraphs; carries
                       a page limit ("shall not exceed two (2) pages"),
                       a stated eval weight ("thirty percent (30%)"),
                       an optional section, and an H3 subsection.
  qform-twin.docx      a buyer questionnaire (attachment-A shape):
                       numbered forms, Question|Response tables with
                       empty answer cells, one pre-answered EXAMPLE row
                       that must NOT become a slot, one 3-column
                       references grid.
  narrative-twin.docx  a narrative core document: prose, a mandated
                       response outline INSIDE the narrative, one
                       FILLED table that must not slot, and the B67-F3
                       fill-in table (role rows present, allocation
                       column empty).

Determinism contract: python-docx re-stamps zip entry timestamps from
the wall clock on save (the xlsx twins' EC-8 disease, same cure) — the
archive is rebuilt with pinned timestamps, stable member order, and
pinned docProps dates. build_* writes the same bytes every run; the
committed goldens are compared byte-for-byte in tests.
"""

import io
import re
import zipfile
from pathlib import Path

from docx import Document

_PINNED_DATE = (2026, 1, 1, 0, 0, 0)
_CORE_DATES = re.compile(
    rb"(<dcterms:(?:created|modified)[^>]*>)[^<]*(</dcterms:(?:created|modified)>)"
)

BUYER = "Northwind Regional Health"


def _freeze(document, path: Path) -> Path:
    """EC-8's cure, docx edition: rebuild the archive pinned."""
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    with zipfile.ZipFile(buffer) as src, zipfile.ZipFile(
        path, "w", zipfile.ZIP_DEFLATED
    ) as dst:
        for name in sorted(src.namelist()):
            data = src.read(name)
            if name == "docProps/core.xml":
                data = _CORE_DATES.sub(
                    rb"\g<1>2026-01-01T00:00:00Z\g<2>", data
                )
            info = zipfile.ZipInfo(name, date_time=_PINNED_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            dst.writestr(info, data)
    return path


def _table(document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            table.rows[r].cells[c].text = text


# ------------------------------------------------------------ template


_TEMPLATE_SECTIONS = [
    "Cover Letter",
    "Executive Summary",
    "Understanding of Your Requirements & Objectives",
    "Proposed Solution & Technical Approach",
    "Implementation Methodology, Project Plan & Timeline",
    "Change Management & Organizational Adoption",
    "Project Team & Key Personnel",
    "Risk Management & Quality Assurance",
    "Company Overview & Qualifications",
    "Relevant Experience, Case Studies & References",
    "Pricing & Commercial Terms",
    "Assumptions, Dependencies & Contractual Terms",
    "Requirements Compliance Matrix",
    "Appendices",
]

_META_ROWS = [
    ["Field", "Value"],
    ["Prepared for (Client)", ""],
    ["RFP title", ""],
    ["RFP / solicitation number", ""],
    ["Submitted by", ""],
    ["Date of submission", ""],
    ["Primary contact", ""],
    ["Due date & method", ""],
]


def build_template_twin(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Firm Response Template")
    doc.add_heading("How to Use This Template", level=1)
    doc.add_paragraph(
        "Complete the metadata table, then work section by section. "
        "Guidance blocks describe what belongs in each section and are "
        "removed from the delivered response."
    )
    _table(doc, _META_ROWS)

    for n, title in enumerate(_TEMPLATE_SECTIONS, start=1):
        doc.add_heading(f"{n}.  {title}", level=1)
        if n == 5:
            doc.add_heading("Timeline & Milestones", level=2)
        guidance = (
            f"▸ WHAT TO INCLUDE — the firm's standard treatment of "
            f"{title.lower()} for this engagement."
        )
        if n == 2:
            guidance += " Keep it to one page."
        _table(doc, [[guidance]])
        if n == 10:
            _table(doc, [[
                "[ Client: [ … ] Scope: [ … ] Outcome: [ … ] ]"
            ]])
        elif n == 11:
            _table(doc, [
                ["Milestone", "Fee", "Duration (weeks)"],
                ["", "", ""],
                ["", "", ""],
            ])
        else:
            _table(doc, [["[ Replace with the drafted section. ]"]])
        if n == 12:
            doc.add_paragraph("Payment schedule & terms: [ … ]")

    return _freeze(doc, path)


# ------------------------------------------------------------- outline


def build_outline_twin(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph(f"Request for Proposal — ERP Implementation Services")
    doc.add_paragraph(f"Issued by {BUYER}.")
    doc.add_heading("Proposal Response Format", level=1)
    doc.add_paragraph(
        "To facilitate the analysis of responses, bidders are required to "
        "prepare their responses in accordance with the instructions "
        "outlined in this section. Bidders must respond in full to all "
        "sections and follow the section numbering in their response. "
        "Failure to follow these instructions may result in rejection."
    )

    doc.add_heading("1. Executive Summary", level=2)
    doc.add_paragraph(
        "Provide a brief narrative describing your organization and the "
        "proposed implementation approach. The executive summary shall "
        "not exceed two (2) pages and should avoid technical jargon."
    )
    doc.add_heading("2. Implementation Approach", level=2)
    doc.add_paragraph(
        "Describe the proposed implementation approach, including phasing "
        "and stages. This section carries thirty percent (30%) of the "
        "evaluation score."
    )
    doc.add_heading("2.1 Project Timeline", level=3)
    doc.add_paragraph(
        "Provide a detailed work plan and schedule in a work breakdown "
        "structure format, including all phases and individual stages."
    )
    doc.add_heading("3. Staffing Plan", level=2)
    doc.add_paragraph(
        "List the proposed project staffing, providing the names and "
        "resumes of key staff and their level of involvement."
    )
    doc.add_heading("4. Change Management", level=2)
    doc.add_paragraph(
        "Clearly identify your approach toward organizational change "
        "management, including any unique approaches or tools."
    )
    doc.add_heading("5. Value-Added Services (Optional)", level=2)
    doc.add_paragraph(
        "This section is optional and may be omitted. Bidders are "
        "encouraged to propose value-added concepts and programs."
    )
    return _freeze(doc, path)


# --------------------------------------------------------------- qform


def build_qform_twin(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Bidder Questionnaire — Attachment A")
    doc.add_paragraph(
        "Bidders are required to answer all questions in the form. "
        "Failure to respond to all questions can result in "
        "disqualification of the entire proposal."
    )
    doc.add_heading("1. Company Background", level=1)
    _table(doc, [
        ["Question", "Response"],
        ["Describe your company's history and ownership structure.", ""],
        ["How many ERP implementations have you completed in the last "
         "five (5) years?", ""],
        ["Do you subcontract any implementation services?", ""],
        ["Describe your quality program.",
         "EXAMPLE: replace this pre-filled row with your own answer."],
    ])
    doc.add_heading("2. Project Management Approach", level=1)
    _table(doc, [
        ["Question", "Response"],
        ["Describe your project governance model.", ""],
        ["Provide your standard escalation path.", ""],
    ])
    doc.add_heading("3. Client References", level=1)
    _table(doc, [
        ["Reference", "Contact", "Phone"],
        ["", "", ""],
        ["", "", ""],
    ])
    return _freeze(doc, path)


# ----------------------------------------------------------- narrative


def build_narrative_twin(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph(
        f"Request for Proposal — ERP Implementation Services for {BUYER}"
    )
    doc.add_heading("Background", level=1)
    doc.add_paragraph(
        f"{BUYER} operates a regional network of care facilities and is "
        "replacing its legacy ERP environment. This solicitation seeks an "
        "implementation partner for the selected platform."
    )
    doc.add_paragraph(
        "The anticipated project phases and their planned quarters are "
        "shown below for reference."
    )
    _table(doc, [
        ["Phase", "Planned quarter"],
        ["Discovery", "Q1"],
        ["Configuration", "Q2"],
        ["Go-live", "Q4"],
    ])
    doc.add_heading("Response Format", level=1)
    doc.add_paragraph(
        "Responses must address the following sections using the section "
        "numbering below."
    )
    doc.add_heading("1. Technical Approach", level=2)
    doc.add_paragraph(
        "Describe the proposed technical approach and architecture."
    )
    doc.add_heading("2. Support Model", level=2)
    doc.add_paragraph(
        "Describe the post-go-live support model, including tiers and "
        "response times."
    )
    doc.add_heading("Resource Commitments", level=1)
    doc.add_paragraph(
        "Bidders shall complete the resource allocation table below."
    )
    _table(doc, [
        ["Role", "Name", "% Allocation"],
        ["Project Manager", "", ""],
        ["Solution Architect", "", ""],
        ["Change Lead", "", ""],
    ])
    return _freeze(doc, path)


# ----------------------------------------------------------------- parts

HEADER_DIRECTIVE = "All responses are due by 5:00 PM local time; late submissions are void."
FOOTER_TEXT = "Confidential — for the named bidder only."
TEXT_BOX_DIRECTIVE = (
    "Note to bidders: describe your data-residency controls in Section 2."
)

_VML_TEXT_BOX = (
    '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml">'
    '<w:pict><v:shape><v:textbox><w:txbxContent><w:p><w:r><w:t>{text}</w:t></w:r>'
    '</w:p></w:txbxContent></v:textbox></v:shape></w:pict></w:r>'
)


def build_parts_twin(path: Path) -> Path:
    """P2-27 (P26b-1, B112): a questionnaire whose buyer instructions sit
    in the parts a body-only walk never reads — a header, a footer and a
    VML text box anchored under a body paragraph (python-docx has no
    text-box API, so the run is raw OOXML). The body itself is the qform
    shape: two numbered sections, one Question|Response table."""
    from docx.oxml import parse_xml

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = HEADER_DIRECTIVE
    doc.sections[0].footer.paragraphs[0].text = FOOTER_TEXT
    doc.add_heading("1. Company Background", level=1)
    anchor = doc.add_paragraph("Answer each question in the table below.")
    anchor._p.append(parse_xml(_VML_TEXT_BOX.format(text=TEXT_BOX_DIRECTIVE)))
    _table(doc, [
        ["Question", "Response"],
        ["Describe your company's history and ownership structure.", ""],
        ["Describe your data-residency controls.", ""],
    ])
    doc.add_heading("2. Project Management Approach", level=1)
    doc.add_paragraph("Describe your project governance model.")
    return _freeze(doc, path)


GOLDENS = {
    "template-twin.docx": build_template_twin,
    "outline-twin.docx": build_outline_twin,
    "qform-twin.docx": build_qform_twin,
    "narrative-twin.docx": build_narrative_twin,
    "parts-twin.docx": build_parts_twin,  # P2-27 (P26b-1)
}


def rebuild_all(directory: Path) -> None:
    for name, builder in GOLDENS.items():
        builder(directory / name)
