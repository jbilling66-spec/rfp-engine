"""P3-15 (P26b-3): output metadata hygiene on every buyer-facing save,
and the part-level report a text-and-style verifier cannot make.

The owner's calls (B119 §1b/§1c): the firm identity is a workspace
`firm.json`, blank when unset and reported as unconfigured, never a
refusal; the firm's own template with tracked changes or comments
refuses the fill by name; the buyer's form written back keeps the
buyer's author and REPORTS its comment parts in the bundle.
"""

import io
import json
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document

from engine.assembly.bundle import compose_bundle
from engine.assembly.docx_writeback import run_docx_writeback
from engine.assembly.hygiene import (firm_identity, inspect_document,
                                     refuse_marked_template, stamp_core_xml)
from engine.assembly.template_fill import (OUTPUT_NAME, WORKING_NAME,
                                           preview_template_fill,
                                           run_template_fill)
from engine.contracts import ContractError, validate
from tests.assembly.test_docx_writeback import _log
from tests.assembly.test_template_fill import (AT, FULL_HAND,
                                               _workspace as _fill_workspace)

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"
FIRM = {"name": "Fixture Advisory LLP", "company": "Fixture Advisory LLP"}


def _set_firm(workspace: Path, firm=FIRM) -> None:
    (Path(workspace) / "firm.json").write_text(json.dumps(firm),
                                                encoding="utf-8")


def _core(path: Path) -> str:
    with zipfile.ZipFile(path) as zf:
        return zf.read("docProps/core.xml").decode("utf-8")


def _with_member(source: Path, target: Path, member: str, data: bytes,
                 *, edit=None) -> Path:
    """Copy a docx adding one member (and optionally editing another) —
    a fixture helper, never a golden."""
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(
            target, "w", zipfile.ZIP_DEFLATED) as dst:
        for name in src.namelist():
            payload = src.read(name)
            if edit and name == edit[0]:
                assert payload.count(edit[1]) >= 1, name
                payload = payload.replace(edit[1], edit[2], 1)
            dst.writestr(name, payload)
        dst.writestr(member, data)
    return target


def _buyer_workspace(tmp_path, source: Path):
    """A client_provided pursuit over THIS buyer file (the
    test_docx_writeback idiom, parameterised on the source bytes so a
    modified twin is the file the slots were parsed from)."""
    import json as _json

    from engine.structure import merge_parsed, parse_buyer_docx
    from engine.workspace import PursuitDir
    from tests.helpers import plant_freeze

    pursuit = PursuitDir(tmp_path, "pur_hyg")
    inbox = pursuit.root / "inbox"
    inbox.mkdir(parents=True, exist_ok=True)
    target = inbox / "qform-twin.docx"
    target.write_bytes(source.read_bytes())
    parsed = parse_buyer_docx(target)
    container = {"pursuit_id": "pur_hyg", **merge_parsed([parsed])}
    pursuit.write_artifact("target_slots", container, name="slots.json")
    plant_freeze(pursuit, "pursuit_plan", {
        "pursuit_id": "pur_hyg", "path": "A_designated",
        "slots_ref": "slots.json", "status": "approved",
        "sections": [{"section_id": "all",
                      "slot_ids": [s["slot_id"] for s in parsed.slots]}]})
    (pursuit.root / "drafts").mkdir(exist_ok=True)
    (pursuit.root / "drafts" / "draft.json").write_text(_json.dumps({
        "plan_sha256": "0" * 64, "revision_n": 1,
        "sections": [{"section_id": "all", "answers": [
            {"slot_id": "s-t00-r01", "status": "drafted",
             "prose": "Founded in 2001, employee-owned."}]}]}),
        encoding="utf-8")
    return pursuit


def _with_creator(source: Path, target: Path, creator: str) -> Path:
    """The qform twin with a BUYER's creator in core.xml (the twin's own
    says python-docx, by construction of the builder)."""
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(
            target, "w", zipfile.ZIP_DEFLATED) as dst:
        for name in src.namelist():
            payload = src.read(name)
            if name == "docProps/core.xml":
                assert payload.count(b"<dc:creator>python-docx</dc:creator>") == 1
                payload = payload.replace(
                    b"<dc:creator>python-docx</dc:creator>",
                    f"<dc:creator>{creator}</dc:creator>".encode())
            dst.writestr(name, payload)
    return target


_COMMENTS_XML = (b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                 b'<w:comments xmlns:w="http://schemas.openxmlformats.org/'
                 b'wordprocessingml/2006/main"/>')
_COMMENTS_REL = (b'<Relationship Id="rIdHygC" Type="http://schemas.openxml'
                 b'formats.org/officeDocument/2006/relationships/comments" '
                 b'Target="comments.xml"/></Relationships>')
_COMMENTS_CT = (b'<Override PartName="/word/comments.xml" ContentType='
                b'"application/vnd.openxmlformats-officedocument.'
                b'wordprocessingml.comments+xml"/></Types>')


def _with_comments(source: Path, target: Path) -> Path:
    """The twin with a comments part the way Word writes one — the part,
    its relationship and its content type — so python-docx carries it
    through a save (an unreferenced member would be dropped)."""
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(
            target, "w", zipfile.ZIP_DEFLATED) as dst:
        for name in src.namelist():
            payload = src.read(name)
            if name == "word/_rels/document.xml.rels":
                payload = payload.replace(b"</Relationships>", _COMMENTS_REL)
            elif name == "[Content_Types].xml":
                payload = payload.replace(b"</Types>", _COMMENTS_CT)
            dst.writestr(name, payload)
        dst.writestr("word/comments.xml", _COMMENTS_XML)
    return target


def test_firm_identity_is_the_workspace_file_blank_when_unset(tmp_path):
    assert firm_identity(tmp_path) == {"name": "", "company": "",
                                       "configured": False}
    _set_firm(tmp_path)
    assert firm_identity(tmp_path) == {**FIRM, "configured": True}
    (tmp_path / "firm.json").write_text("{", encoding="utf-8")
    with pytest.raises(ContractError, match="not readable JSON"):
        firm_identity(tmp_path)


def test_a_buyer_form_written_back_keeps_the_buyers_author(tmp_path):
    """Buyer-owned: creator and created stand; the firm is the last
    modifier; no generator string anywhere in core.xml."""
    buyer_file = _with_creator(FIXTURES / "qform-twin.docx",
                               tmp_path / "buyer.docx", "Buyer Procurement")
    pursuit = _buyer_workspace(tmp_path, buyer_file)
    _set_firm(tmp_path)
    source = pursuit.root / "inbox" / "qform-twin.docx"
    before = Document(str(source)).core_properties
    assert before.author == "Buyer Procurement"
    log = _log(pursuit)
    facts = run_docx_writeback(pursuit, log, at=AT, confirmed_by="Pat Lead")
    log.run_end(status="completed")
    out = Document(str(pursuit.root / facts["output_file"])).core_properties
    assert out.author == "Buyer Procurement"
    assert out.created == before.created
    assert out.last_modified_by == FIRM["name"]
    assert out.modified.isoformat().startswith(AT[:19])
    core = _core(pursuit.root / facts["output_file"]).lower()
    assert "python-docx" not in core and "generated by" not in core


def test_the_bundle_reports_a_buyer_forms_comment_parts(tmp_path):
    """Buyer-owned files REPORT: a questionnaire carrying the buyer's own
    comments part writes back and the bundle names the part."""
    marked = _with_comments(FIXTURES / "qform-twin.docx", tmp_path / "m.docx")
    assert inspect_document(marked)["comment_parts"] == ["word/comments.xml"]
    pursuit = _buyer_workspace(tmp_path, marked)
    _set_firm(tmp_path)
    log = _log(pursuit)
    run_docx_writeback(pursuit, log, at=AT, confirmed_by="Pat Lead")
    bundle = compose_bundle(pursuit, log, at=AT, composed_by="Pat Lead")
    log.run_end(status="completed")
    validate("submission_bundle", bundle)
    entry = next(d for d in bundle["deliverables"]
                 if d["lane"] == "docx_writeback")
    assert entry["status"] == "produced"
    assert entry["hygiene"]["comment_parts"] == ["word/comments.xml"]
    assert entry["hygiene"]["revision_marks"] == {}
    assert entry["hygiene"]["generator_strings"] == []
    assert entry["hygiene"]["last_modified_by"] == FIRM["name"]
    assert entry["hygiene"]["firm_identity"] == "configured"


def test_an_unset_firm_stamps_blank_and_reports_unconfigured(tmp_path):
    """Never a refusal: the output carries NO author rather than the
    generator's, and the bundle says the identity was unconfigured."""
    pursuit = _buyer_workspace(tmp_path, FIXTURES / "qform-twin.docx")
    log = _log(pursuit)
    facts = run_docx_writeback(pursuit, log, at=AT, confirmed_by="Pat Lead")
    bundle = compose_bundle(pursuit, log, at=AT, composed_by="Pat Lead")
    log.run_end(status="completed")
    out = Document(str(pursuit.root / facts["output_file"])).core_properties
    assert out.last_modified_by == ""
    # the twin's own creator names the generator — stripped, never kept
    assert out.author == ""
    entry = next(d for d in bundle["deliverables"]
                 if d["lane"] == "docx_writeback")
    assert entry["hygiene"]["firm_identity"] == "unconfigured"
    assert entry["hygiene"]["generator_strings"] == []


def test_the_template_fill_is_firm_owned_in_both_copies(tmp_path):
    """Firm-owned: author AND last-modified-by are the firm, the
    descriptive fields cleared — the bundled template's own creator
    (python-docx, by construction of the twin) never reaches anyone."""
    pursuit = _fill_workspace(tmp_path, all_prose=True, hand=FULL_HAND)
    _set_firm(tmp_path)
    log = _log(pursuit)
    facts = run_template_fill(pursuit, log, confirmed_by="Pat Lead", at=AT)
    log.run_end(status="completed")
    assert facts["buyer_copy_produced"] is True
    for name in (WORKING_NAME, OUTPUT_NAME):
        props = Document(str(pursuit.root / name)).core_properties
        assert props.author == FIRM["name"], name
        assert props.last_modified_by == FIRM["name"], name
        assert props.comments == "" and props.subject == ""
        core = _core(pursuit.root / name).lower()
        assert "python-docx" not in core and "generated by" not in core
    assert Document(str(pursuit.root / WORKING_NAME)).core_properties.title \
        .endswith("(working copy)")


def test_a_marked_template_refuses_the_fill_by_name(tmp_path, monkeypatch):
    """The firm-owned door: a template carrying a tracked insertion and a
    comments part refuses preview AND confirm, naming the counts; no
    copy is produced."""
    import engine.assembly.template_fill as fill_mod
    from engine.planning.plan import REFERENCE_DEFAULT

    marked = _with_member(
        REFERENCE_DEFAULT, tmp_path / "marked-template.docx",
        "word/comments.xml",
        b'<?xml version="1.0" encoding="UTF-8"?><w:comments xmlns:w="http://'
        b'schemas.openxmlformats.org/wordprocessingml/2006/main"/>',
        edit=("word/document.xml", b"<w:r>",
              b'<w:ins w:id="9" w:author="Reviewer" w:date="2026-01-01T00:00:00Z"><w:r>'))
    # the insertion must close: wrap the FIRST run only
    data = bytearray(marked.read_bytes())
    with zipfile.ZipFile(io.BytesIO(bytes(data))) as zf:
        document = zf.read("word/document.xml")
    assert document.count(b"<w:ins ") == 1
    closed = document.replace(b"</w:r>", b"</w:r></w:ins>", 1)
    marked = _with_member(marked, tmp_path / "marked2.docx",
                          "word/_hygiene_marker.txt", b"x",
                          edit=("word/document.xml", document, closed))
    report = inspect_document(marked)
    assert report["revision_marks"] == {"w:ins": 1}
    assert report["comment_parts"] == ["word/comments.xml"]
    with pytest.raises(ContractError, match=r"1 tracked-change mark"):
        refuse_marked_template(marked)

    import hashlib
    sha = hashlib.sha256(marked.read_bytes()).hexdigest()
    pursuit = _fill_workspace(tmp_path, ref_sha=sha, all_prose=True,
                              hand=FULL_HAND)
    monkeypatch.setattr(fill_mod, "REFERENCE_DEFAULT", marked)
    with pytest.raises(ContractError, match="word/comments.xml"):
        preview_template_fill(pursuit, at=AT)
    log = _log(pursuit)
    with pytest.raises(ContractError, match="tracked-change"):
        run_template_fill(pursuit, log, confirmed_by="Pat Lead", at=AT)
    log.run_end(status="failed")
    assert not (pursuit.root / WORKING_NAME).exists()
    assert not (pursuit.root / OUTPUT_NAME).exists()


def test_stamp_core_xml_applies_the_buyer_owned_rule():
    """The raw-part twin of stamp_docx, for the xlsx zip patch (step 6):
    creator kept unless it names a generator, the firm as last modifier,
    modified = at."""
    core = (b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/'
            b'package/2006/metadata/core-properties" xmlns:dc="http://purl.org/'
            b'dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" '
            b'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
            b'<dc:creator>Buyer Procurement</dc:creator>'
            b'<dc:description>generated by openpyxl</dc:description>'
            b'<cp:lastModifiedBy>Someone</cp:lastModifiedBy>'
            b'<dcterms:modified xsi:type="dcterms:W3CDTF">2020-01-01T00:00:00Z'
            b'</dcterms:modified></cp:coreProperties>')
    out = stamp_core_xml(core, firm={**FIRM, "configured": True},
                         at="2026-09-04T10:00:00Z").decode("utf-8")
    assert "<dc:creator>Buyer Procurement</dc:creator>" in out
    assert "<dc:description/>" in out or "<dc:description></dc:description>" in out
    assert f"<cp:lastModifiedBy>{FIRM['name']}</cp:lastModifiedBy>" in out
    assert re.search(r"<dcterms:modified[^>]*>2026-09-04T10:00:00Z<", out)


def test_the_rendered_submission_and_review_are_firm_owned(tmp_path):
    """The two bare-Document lanes: python-docx's default core.xml is
    replaced with the firm's, the title from the frozen brief."""
    from engine.assembly.docx import (REVIEW_NAME, SUBMISSION_NAME,
                                      render_review, render_submission)
    from tests.validation.fixtures.validations import run_validation_package

    pursuit, report, _ = run_validation_package(tmp_path)
    assert report.status == "complete"
    _set_firm(tmp_path)
    log = _log(pursuit)
    render_submission(pursuit, log, at=AT)
    render_review(pursuit, log, at=AT)
    log.run_end(status="completed")
    for name in (SUBMISSION_NAME, REVIEW_NAME):
        props = Document(str(pursuit.root / name)).core_properties
        assert props.author == FIRM["name"], name
        assert props.last_modified_by == FIRM["name"], name
        assert props.title.startswith("Response"), name
        assert props.comments == ""
        core = _core(pursuit.root / name).lower()
        assert "python-docx" not in core and "generated by" not in core
    review = Document(str(pursuit.root / REVIEW_NAME)).core_properties
    assert review.title.endswith("(internal review)")
