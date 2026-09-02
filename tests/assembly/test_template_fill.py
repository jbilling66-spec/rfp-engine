"""In-place firm-template fill acceptance (P17/C9, B73§2 closed; P26a
item 1, P1-27): replacement-not-insertion — drafted section prose takes
the guidance box's place AND the placeholder table goes too; the
template's own authoring scaffolding (title, "How to Use", instruction)
leaves EVERY copy; hand-completion values render into the metadata
record, the pricing grid, the case block, and the inline line; a section
a human still drafts keeps its guidance IDENTICAL in the working copy;
the buyer copy is written only when nothing remains — proven on the
bundled template; the stream-diff verifier refuses any drift outside
the declared change set; the template is only ever filled at the exact
digest the frozen plan was built against.

Workspaces are hand-built through the real seams (the P16
tests/assembly precedent): slots via the real parser + merge, the plan
and envelope as raw artifacts the engine reads without re-validating,
hand values through the hand-fill writer.
"""

import hashlib
import json
import shutil

import pytest
from docx import Document

from engine.assembly.hand_fill import write_hand_fill
from engine.assembly.template_fill import (
    OUTPUT_NAME,
    WORKING_NAME,
    _assert_fill_roundtrip,
    _body_stream,
    preview_template_fill,
    run_template_fill,
    withheld_reason,
)
from engine.contracts import ContractError
from engine.llm import effective_config
from engine.planning.plan import REFERENCE_DEFAULT
from engine.runlog import RunLogger, read_run
from engine.structure import merge_parsed, parse_default_template
from engine.version import engine_version
from engine.workspace import PursuitDir
from tests.helpers import plant_freeze

AT = "2026-08-29T12:00:00Z"
PARA_ONE = "First paragraph of the summary."
PARA_TWO = "Second paragraph, still synthetic."
TEMPLATE_SHA = hashlib.sha256(REFERENCE_DEFAULT.read_bytes()).hexdigest()
# the twelve template sections that carry a ▸ prose slot (10 and 11 are
# the case block and the pricing grid — hand shapes, no prose)
PROSE_SECTIONS = [1, 2, 3, 4, 5, 6, 7, 8, 9, 12, 13, 14]
SCAFFOLDING = ["Firm Response Template", "How to Use This Template",
               "Complete the metadata table, then work section by section. "
               "Guidance blocks describe what belongs in each section and "
               "are removed from the delivered response."]
META = {"prepared_for_client": "Synthetic Buyer Co",
        "rfp_title": "Synthetic Modernisation RFP",
        "rfp_solicitation_number": "RFP-0001", "submitted_by": "The Firm",
        "date_of_submission": "2026-09-30", "primary_contact": "Pat Lead",
        "due_date_method": "2026-10-01, portal"}
GRID = [{"milestone": "Kickoff", "fee": "$1,000", "duration_weeks": "2"},
        {"milestone": "Build", "fee": "$4,500", "duration_weeks": "6"},
        {"milestone": "Handover", "fee": "$500", "duration_weeks": "1"}]
CASES = [{"client": "A synthetic utility", "scope": "Finance system",
          "outcome": "Live on schedule"},
         {"client": "A synthetic college", "scope": "Payroll",
          "outcome": "Zero defects at go-live"}]
INLINE = "Net 30 from invoice; 40% on kickoff"
FULL_HAND = {"s-front-meta": META, "s-h11": GRID, "s-h10": CASES,
             "s-h12-1": INLINE}


def _workspace(tmp_path, *, ref_sha=None, source_mode=None,
               all_prose=False, hand=None):
    pursuit = PursuitDir(tmp_path, "pur_fill")
    parsed = parse_default_template(REFERENCE_DEFAULT)
    container = {"pursuit_id": "pur_fill", **merge_parsed([parsed])}
    if source_mode:
        container["source_mode"] = source_mode
        (pursuit.root / "slots.json").write_text(json.dumps(container))
    else:
        pursuit.write_artifact("target_slots", container, name="slots.json")
    pursuit.checkpoint("path_b_outline", {
        "reference_sha256": ref_sha or TEMPLATE_SHA})
    if all_prose:
        ids = [f"s-h{n:02d}" for n in PROSE_SECTIONS]
        drafted = {s: f"Synthetic prose for template section {s}."
                   for s in ids}
    else:
        ids = ["s-h02", "s-h03", "s-h04"]
        drafted = {"s-h02": f"{PARA_ONE}\n\n{PARA_TWO}",
                   "s-h03": "One firm paragraph."}
    sections = [{"section_id": f"sec-{slot}",
                 "slot_ids": [f"{slot}-hdr", slot]} for slot in ids]
    plant_freeze(pursuit, "pursuit_plan", {
        "pursuit_id": "pur_fill", "path": "B_free_flow",
        "slots_ref": "slots.json", "status": "approved",
        "sections": sections})
    envelope = [{"section_id": f"sec-{s}", "status": "drafted",
                 "prose": drafted[s]} if s in drafted
                else {"section_id": f"sec-{s}",
                      "status": "awaiting_disposition"} for s in ids]
    (pursuit.root / "drafts" / "draft.json").write_text(json.dumps({
        "plan_sha256": "0" * 64, "revision_n": 0, "sections": envelope}))
    if hand:
        write_hand_fill(pursuit, container=container,
                        template_sha256=TEMPLATE_SHA, entered_by="Pat Lead",
                        at=AT, values=hand)
    return pursuit


def _log(pursuit):
    log = RunLogger(pursuit.root, pursuit.new_run_id(), pursuit.pursuit_id)
    log.run_start(mode="dry_run", engine_version=engine_version(),
                  config=effective_config(), kb_snapshot="kb@empty")
    return log


def _texts(doc) -> list[str]:
    """Every paragraph AND every cell text, for scaffolding sweeps."""
    out = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            out.extend(cell.text for cell in row.cells)
    return out


def _firsts(doc) -> list[str]:
    return [t.rows[0].cells[0].text.strip() for t in doc.tables]


def test_fill_replaces_guidance_removes_placeholder_and_reports(tmp_path):
    """THE acceptance arc: filled sections lose guidance AND placeholder,
    the delivered paragraphs land, everything else survives identical in
    the WORKING copy, the facts record every template slot exactly once
    and itemize what a human still owes, and the buyer copy is withheld."""
    pursuit = _workspace(tmp_path)
    facts = preview_template_fill(pursuit, at=AT)
    assert facts["confirmed_by"] == "(unconfirmed)"
    decisions = {r["slot_id"]: r["decision"] for r in facts["sections"]}
    assert decisions["s-h02"] == "filled"
    assert decisions["s-h03"] == "filled"
    assert decisions["s-h04"] == "kept_guidance"
    assert decisions["s-h01"] == "refused_unnamed"
    assert {decisions[s] for s in
            ("s-front-meta", "s-h10", "s-h11", "s-h12-1")} \
        == {"fill_by_hand"}
    filled_rows = [r for r in facts["sections"]
                   if r["decision"] == "filled"]
    assert all(r["placeholder_removed"] for r in filled_rows)
    assert next(r for r in filled_rows
                if r["slot_id"] == "s-h02")["paragraphs"] == 2
    assert facts["buyer_copy_produced"] is False

    log = _log(pursuit)
    out = run_template_fill(pursuit, log, confirmed_by="Pat Lead", at=AT)
    log.run_end(status="completed")
    assert out["confirmed_by"] == "Pat Lead"
    assert not (pursuit.root / OUTPUT_NAME).exists(), \
        "the buyer copy is withheld while anything remains"

    doc = Document(str(pursuit.root / WORKING_NAME))
    texts = [p.text for p in doc.paragraphs]
    assert PARA_ONE in texts and PARA_TWO in texts
    assert texts.index(PARA_TWO) == texts.index(PARA_ONE) + 1, \
        "multi-paragraph prose stays multi-paragraph, in order"
    firsts = _firsts(doc)
    assert len([t for t in firsts if t.startswith("▸")]) == 12  # 14 − 2
    assert len([t for t in firsts if t.startswith("[ Replace")]) == 10
    assert len(out["remaining_guidance"]) == 10  # 12 prose − 2 filled
    assert not any("executive summary" in g.lower()
                   for g in out["remaining_guidance"])
    assert len(out["remaining_by_hand"]) == 4
    assert any(line.startswith("Pricing & Commercial Terms: missing")
               for line in out["remaining_by_hand"])

    # The template at config/ is never mutated, and the output reparses.
    assert TEMPLATE_SHA == out["template_sha256"]
    reparsed = parse_default_template(pursuit.root / WORKING_NAME)
    assert reparsed.slot_count == 29  # 31 − the two consumed prose slots

    # The facts artifact and the working copy are on the run log.
    artifacts = [r for r in read_run(pursuit.root / "runs" / log.run_id /
                                     "run.jsonl")
                 if r["record_type"] == "artifact"]
    kinds = {a["artifact"]["kind"] for a in artifacts}
    assert {"template_fill_facts", "export"} <= kinds
    assert (pursuit.root / "exports" / "template-fill-facts.json").is_file()


def test_fill_strips_the_front_matter_and_keeps_unfilled_guidance_in_the_working_copy(
        tmp_path):
    """P1-27's first half: the title, the "How to Use This Template"
    heading, and the instruction paragraph leave every copy; the
    metadata table stays (it is the cover block); a section nobody
    drafted keeps its guidance and placeholder for the hand-drafter."""
    pursuit = _workspace(tmp_path)
    out = run_template_fill(pursuit, _log(pursuit), confirmed_by="Pat",
                            at=AT)
    assert out["scaffolding_removed"] == SCAFFOLDING
    doc = Document(str(pursuit.root / WORKING_NAME))
    stream = _body_stream(doc)
    assert stream[0][0] == "t" and stream[0][1][0][0] == "Field", \
        "the working copy opens on the metadata table — nothing before it"
    joined = "\n".join(_texts(doc))
    for phrase in ("How to Use", "Firm Response Template",
                   "removed from the delivered response"):
        assert phrase not in joined
    kept = next(s["question_text"]
                for s in parse_default_template(REFERENCE_DEFAULT).slots
                if s["slot_id"] == "s-h04")  # awaiting_disposition above
    assert kept.startswith("▸") and kept in _firsts(doc), \
        "an undrafted section's guidance survives identical"


def test_hand_fill_values_land_in_record_grid_case_block_and_inline(tmp_path):
    """The hand-completion record renders in place: metadata values into
    column 1 of the record, one grid row per entry (rows ADDED past the
    template's two blanks), the case block as labeled lines per entry,
    the inline line's bracket span replaced — and the guidance box
    above a hand-completed table goes with it."""
    pursuit = _workspace(tmp_path, hand=FULL_HAND)
    out = run_template_fill(pursuit, _log(pursuit), confirmed_by="Pat",
                            at=AT)
    decisions = {r["slot_id"]: r for r in out["sections"]}
    for slot_id in FULL_HAND:
        assert decisions[slot_id]["decision"] == "filled_by_hand", slot_id
    assert decisions["s-h11"]["fields_written"] \
        == ["duration_weeks", "fee", "milestone"]
    assert out["remaining_by_hand"] == []

    doc = Document(str(pursuit.root / WORKING_NAME))
    meta = doc.tables[0]
    assert meta.rows[0].cells[0].text == "Field"  # header kept: working copy
    assert [(r.cells[0].text, r.cells[1].text) for r in meta.rows[1:]] == [
        ("Prepared for (Client)", META["prepared_for_client"]),
        ("RFP title", META["rfp_title"]),
        ("RFP / solicitation number", META["rfp_solicitation_number"]),
        ("Submitted by", META["submitted_by"]),
        ("Date of submission", META["date_of_submission"]),
        ("Primary contact", META["primary_contact"]),
        ("Due date & method", META["due_date_method"])]
    grid = next(t for t in doc.tables
                if t.rows[0].cells[0].text == "Milestone")
    assert [[c.text for c in r.cells] for r in grid.rows] == [
        ["Milestone", "Fee", "Duration (weeks)"],
        ["Kickoff", "$1,000", "2"], ["Build", "$4,500", "6"],
        ["Handover", "$500", "1"]]
    case = next(t for t in doc.tables
                if t.rows[0].cells[0].text.startswith("Client: A synthetic"))
    assert case.rows[0].cells[0].text == (
        "Client: A synthetic utility\nScope: Finance system\n"
        "Outcome: Live on schedule\n\nClient: A synthetic college\n"
        "Scope: Payroll\nOutcome: Zero defects at go-live")
    assert f"Payment schedule & terms: {INLINE}" in \
        [p.text for p in doc.paragraphs]
    firsts = _firsts(doc)
    assert len([t for t in firsts if t.startswith("▸")]) == 10  # 12 − §10/§11
    assert not any(t.startswith("[") for t in firsts if "Client" in t)


def test_hand_fill_record_from_another_template_refuses(tmp_path):
    pursuit = _workspace(tmp_path)
    container = pursuit.read_artifact("slots.json")
    write_hand_fill(pursuit, container=container, template_sha256="b" * 64,
                    entered_by="Pat", at=AT, values={"s-h12-1": INLINE})
    with pytest.raises(ContractError, match="different firm template"):
        preview_template_fill(pursuit, at=AT)


def test_buyer_copy_is_withheld_while_anything_remains(tmp_path):
    """Every prose section drafted, two hand slots still owed: the buyer
    copy is not written (a stale one is removed), the facts itemize the
    two, and the bundle reason names them and the working copy."""
    pursuit = _workspace(tmp_path, all_prose=True,
                         hand={"s-front-meta": META, "s-h12-1": INLINE})
    stale = pursuit.root / OUTPUT_NAME
    stale.parent.mkdir(parents=True, exist_ok=True)
    stale.write_bytes(b"stale bytes from an earlier fill")
    out = run_template_fill(pursuit, _log(pursuit), confirmed_by="Pat",
                            at=AT)
    assert out["remaining_guidance"] == []
    assert [line.split(":")[0] for line in out["remaining_by_hand"]] == [
        "Relevant Experience, Case Studies & References",
        "Pricing & Commercial Terms"]
    assert out["buyer_copy_produced"] is False
    assert not stale.exists()
    assert (pursuit.root / WORKING_NAME).is_file()
    reason = withheld_reason(out)
    assert reason.startswith("buyer copy withheld — 2 item(s) remain")
    assert "Pricing & Commercial Terms: missing at least one row" in reason
    assert reason.endswith(f"working copy at {WORKING_NAME}")


def test_buyer_copy_from_the_bundled_template_has_no_scaffolding(tmp_path):
    """THE P1-27 acceptance clause, on the bundled template: every prose
    section drafted + the hand-completion record complete → the buyer
    copy is produced and carries no authoring scaffolding at all — no
    guidance marker, no bracket, no "How to Use", no "Field" header —
    while every drafted paragraph and every hand value is present."""
    pursuit = _workspace(tmp_path, all_prose=True, hand=FULL_HAND)
    out = run_template_fill(pursuit, _log(pursuit), confirmed_by="Pat",
                            at=AT)
    assert out["buyer_copy_produced"] is True
    assert out["remaining_guidance"] == [] and out["remaining_by_hand"] == []
    assert (pursuit.root / OUTPUT_NAME).is_file()
    assert (pursuit.root / WORKING_NAME).is_file()

    doc = Document(str(pursuit.root / OUTPUT_NAME))
    texts = _texts(doc)
    joined = "\n".join(texts)
    assert "▸" not in joined and "[" not in joined and "]" not in joined
    for phrase in ("How to Use", "Firm Response Template", "Replace with",
                   "removed from the delivered response"):
        assert phrase not in joined
    assert "Field" not in texts, "the cover block never says 'Field'"
    assert doc.tables[0].rows[0].cells[0].text == "Prepared for (Client)"
    for n in PROSE_SECTIONS:
        assert f"Synthetic prose for template section s-h{n:02d}." in texts
    for value in list(META.values()) + [INLINE]:
        assert any(value in t for t in texts), value
    grid = next(t for t in doc.tables
                if t.rows[0].cells[0].text == "Milestone")
    assert len(grid.rows) == 4 and grid.rows[3].cells[0].text == "Handover"
    assert not any(all(c.text == "" for c in r.cells) for r in grid.rows), \
        "no blank template rows reach the buyer"
    assert "Outcome: Zero defects at go-live" in joined


def test_verifier_refuses_an_off_target_paragraph_or_cell_mutation(tmp_path):
    """The negative proof: a document differing anywhere outside the
    declared change set is refused, never handed back — a paragraph
    edit and a cell edit alike."""
    pursuit = _workspace(tmp_path)
    tampered = pursuit.root / "exports" / "submission" / "tampered.docx"
    tampered.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE_DEFAULT, tampered)
    doc = Document(str(tampered))
    victim = next(p for p in doc.paragraphs if p.text.strip())
    victim.text = victim.text + " [drifted]"
    doc.save(str(tampered))
    with pytest.raises(ContractError, match="drifted outside"):
        _assert_fill_roundtrip(REFERENCE_DEFAULT, tampered, {})
    shutil.copy2(REFERENCE_DEFAULT, tampered)
    doc = Document(str(tampered))
    doc.tables[0].rows[1].cells[1].text = "smuggled"
    doc.save(str(tampered))
    with pytest.raises(ContractError, match="drifted outside"):
        _assert_fill_roundtrip(REFERENCE_DEFAULT, tampered, {})


def test_template_drift_since_planning_refuses(tmp_path):
    pursuit = _workspace(tmp_path, ref_sha="0" * 64)
    with pytest.raises(ContractError, match="drifted since planning"):
        preview_template_fill(pursuit, at=AT)


def test_buyer_containers_refuse_this_lane(tmp_path):
    pursuit = _workspace(tmp_path, source_mode="client_provided")
    with pytest.raises(ContractError, match="firm_default lane"):
        preview_template_fill(pursuit, at=AT)
