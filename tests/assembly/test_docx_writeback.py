"""DOCX write-back acceptance (P16/C8, tests/assembly is born here):
the buyer's questionnaire is filled in a COPY, only plan-named drafted
prose lands, every non-write is a recorded refusal, and THE named test
proves buyer content undisturbed via the document-model diff.

Workspaces are hand-built through the real seams (slots via the real
parser + container merge; plan/envelope as raw artifacts the engine
reads without re-validating — the refusal-suite precedent).
"""

import json
import shutil
from pathlib import Path

import pytest
from docx import Document

from engine.assembly.docx_writeback import (
    compute_docx_facts,
    preview_docx_writeback,
    run_docx_writeback,
)
from engine.contracts import ContractError, validate
from engine.llm import effective_config
from engine.runlog import RunLogger
from engine.structure import (
    merge_parsed,
    parse_buyer_docx,
    parse_default_template,
)
from engine.version import engine_version
from engine.workspace import PursuitDir

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"
AT = "2026-08-28T12:00:00Z"

PROSE = "Founded in 2001, employee-owned, ERP delivery is the practice."
GOV = None  # planned but undrafted — the honest empty


def _workspace(tmp_path, source_name: str, *, template=False):
    pursuit = PursuitDir(tmp_path, "pur_docxwb")
    inbox = pursuit.root / "inbox"
    inbox.mkdir(parents=True, exist_ok=True)
    if template:
        parsed = parse_default_template(FIXTURES / "template-twin.docx")
    else:
        shutil.copy2(FIXTURES / source_name, inbox / source_name)
        parsed = parse_buyer_docx(inbox / source_name)
    container = {"pursuit_id": "pur_docxwb", **merge_parsed([parsed])}
    pursuit.write_artifact("target_slots", container, name="slots.json")

    planned = [s["slot_id"] for s in parsed.slots
               if s["slot_id"] != "s-t01-r02"]  # one slot left UNplanned
    (pursuit.root / "plan.frozen.json").write_text(json.dumps({
        "pursuit_id": "pur_docxwb", "path": "A_designated",
        "slots_ref": "slots.json", "status": "approved",
        "sections": [{"section_id": "all", "slot_ids": planned}],
    }), encoding="utf-8")
    answers = [{"slot_id": "s-t00-r01", "status": "drafted", "prose": PROSE},
               {"slot_id": "s-t01-r01", "status": "awaiting_disposition"}]
    (pursuit.root / "drafts").mkdir(exist_ok=True)
    (pursuit.root / "drafts" / "draft.json").write_text(json.dumps({
        "plan_sha256": "0" * 64, "revision_n": 1,
        "sections": [{"section_id": "all", "answers": answers}],
    }), encoding="utf-8")
    return pursuit


def _log(pursuit):
    log = RunLogger(pursuit.root, pursuit.new_run_id(), pursuit.pursuit_id)
    log.run_start(mode="dry_run", engine_version=engine_version(),
                  config=effective_config(), kb_snapshot="kb@empty")
    return log


@pytest.fixture()
def qform(tmp_path):
    return _workspace(tmp_path, "qform-twin.docx")


def test_preview_records_every_decision(qform):
    facts = preview_docx_writeback(qform, at=AT)
    validate("writeback_facts", facts)
    by_slot = {c["slot_id"]: c for c in facts["cells"]}
    assert by_slot["s-t00-r01"]["decision"] == "written"
    assert by_slot["s-t00-r01"]["before"] == ""
    assert by_slot["s-t00-r01"]["after"] == PROSE
    assert by_slot["s-t00-r02"]["decision"] == "refused_shape"  # numeric
    assert by_slot["s-t00-r03"]["decision"] == "refused_shape"  # boolean
    assert by_slot["s-t01-r01"]["decision"] == "empty_no_prose"
    assert by_slot["s-t01-r02"]["decision"] == "refused_unnamed"
    assert by_slot["s-t02"]["decision"] == "refused_shape"  # the grid
    assert facts["confirmed_by"] == "(unconfirmed)"
    # the preview wrote NOTHING (the workspace scaffold may hold the
    # empty directory; what matters is no file)
    exports = qform.root / "exports"
    assert not exports.exists() or not any(exports.rglob("*"))


def test_run_fills_the_copy_and_only_the_copy(qform):
    log = _log(qform)
    facts = run_docx_writeback(qform, log, at=AT, confirmed_by="Pat Lead")
    log.run_end(status="completed")
    output = qform.root / facts["output_file"]
    assert output.exists()
    doc = Document(str(output))
    cell = doc.tables[0].rows[1].cells[1]
    assert cell.text == PROSE
    # the buyer's pre-filled EXAMPLE row is untouched
    assert doc.tables[0].rows[4].cells[1].text.startswith("EXAMPLE:")
    # the inbox original never changed
    original = Document(str(qform.root / facts["source_file"]))
    assert original.tables[0].rows[1].cells[1].text == ""
    validate("writeback_facts",
             json.loads((qform.root / "exports"
                         / "docx-writeback-facts.json").read_text()))


def test_buyer_content_is_undisturbed(qform):
    """THE named acceptance test: outside the single written cell, the
    document MODEL of the copy equals the buyer's original — every
    paragraph (text and style), every table cell."""
    log = _log(qform)
    facts = run_docx_writeback(qform, log, at=AT, confirmed_by="Pat Lead")
    log.run_end(status="completed")
    src = Document(str(qform.root / facts["source_file"]))
    out = Document(str(qform.root / facts["output_file"]))
    assert ([(p.text, p.style.name) for p in src.paragraphs]
            == [(p.text, p.style.name) for p in out.paragraphs])
    written = {(c["table_index"], c["row"], c["column"])
               for c in facts["cells"] if c["decision"] == "written"}
    assert written == {(0, 1, 1)}
    for t, (st, ot) in enumerate(zip(src.tables, out.tables)):
        for r, (sr, orow) in enumerate(zip(st.rows, ot.rows)):
            for c, (sc, oc) in enumerate(zip(sr.cells, orow.cells)):
                if (t, r, c) not in written:
                    assert sc.text == oc.text, (t, r, c)


def test_outline_sections_refuse_no_cell_never_silently(tmp_path):
    pursuit = _workspace(tmp_path, "outline-twin.docx")
    facts = preview_docx_writeback(pursuit, at=AT)
    decisions = {c["decision"] for c in facts["cells"]
                 if c["slot_id"].startswith("s-r")}
    assert "refused_no_cell" in decisions
    no_cell = next(c for c in facts["cells"]
                   if c["decision"] == "refused_no_cell")
    assert "rendered response document" in no_cell["reason"]
    assert no_cell["docx_anchor"]  # the anchor still names the section


def test_firm_template_container_routes_to_the_fill_lane(tmp_path):
    """P16 recorded this refusal as 'the P17 lane'; P17/C10 made the
    lane real — the buyer-docx engine still refuses firm_default, now
    pointing at template_fill (the web dispatcher routes there by
    source_mode)."""
    pursuit = _workspace(tmp_path, "qform-twin.docx", template=True)
    with pytest.raises(ContractError, match="template_fill"):
        compute_docx_facts(pursuit, at=AT, confirmed_by="x")


def test_wrong_source_refuses(qform):
    (qform.root / "inbox" / "qform-twin.docx").unlink()
    shutil.copy2(FIXTURES / "outline-twin.docx",
                 qform.root / "inbox" / "outline-twin.docx")
    with pytest.raises(ContractError, match="EXACT file"):
        preview_docx_writeback(qform, at=AT)


# -- P18/C5: multi-docx write-back lands (B74§3g retires) ----------------

@pytest.fixture()
def two_docx(tmp_path):
    """qform-twin (f00, a table questionnaire) + outline-twin (f01, a
    mandated outline with NO answer tables) declared together — the
    exact multi-form shape B74§3g refused."""
    from engine.assembly.bundle import declared_deliverables

    pursuit = PursuitDir(tmp_path, "pur_docxwb")
    inbox = pursuit.root / "inbox"
    for name in ("qform-twin.docx", "outline-twin.docx"):
        shutil.copy2(FIXTURES / name, inbox / name)
    parsed = [parse_buyer_docx(inbox / "qform-twin.docx"),
              parse_buyer_docx(inbox / "outline-twin.docx")]
    container = {"pursuit_id": "pur_docxwb", **merge_parsed(parsed)}
    pursuit.write_artifact("target_slots", container, name="slots.json")
    planned = [s["slot_id"] for s in container["slots"]
               if not s.get("is_header")]
    (pursuit.root / "plan.frozen.json").write_text(json.dumps({
        "pursuit_id": "pur_docxwb", "path": "A_designated",
        "slots_ref": "slots.json", "status": "approved",
        "sections": [{"section_id": "all", "slot_ids": planned}],
    }), encoding="utf-8")
    (pursuit.root / "drafts").mkdir(exist_ok=True)
    (pursuit.root / "drafts" / "draft.json").write_text(json.dumps({
        "plan_sha256": "0" * 64, "revision_n": 1,
        "sections": [{"section_id": "all", "answers": [
            {"slot_id": "f00-s-t00-r01", "status": "drafted",
             "prose": PROSE}]}],
    }), encoding="utf-8")
    bindings = declared_deliverables(pursuit,
                                     pursuit.read_artifact("slots.json"))
    return pursuit, bindings


def test_two_docx_sources_each_get_their_own_filled_copy(two_docx):
    pursuit, bindings = two_docx
    assert [b["prefix"] for b in bindings] == ["f00-", "f01-"]
    for binding in bindings:
        log = _log(pursuit)
        run_docx_writeback(pursuit, log, at=AT, confirmed_by="Pat Lead",
                           binding=binding)
        log.run_end(status="completed")
    qform_out = Document(str(
        pursuit.root / "exports" / "writeback" / "qform-twin.docx"))
    assert qform_out.tables[0].rows[1].cells[1].text == PROSE
    assert (pursuit.root / "exports" / "writeback"
            / "outline-twin.docx").exists()
    for facts_name in ("exports/docx-writeback-facts-f00.json",
                       "exports/docx-writeback-facts-f01.json"):
        validate("writeback_facts", pursuit.read_artifact(facts_name))


def test_union_of_per_file_records_accounts_for_every_slot(two_docx):
    """THE named coverage test (B77§2 D7): the cross-file filter is a
    per-binding contract, not a silent drop — the UNION of the per-file
    cells[] covers every non-header slot of the container, so no slot
    of any declared form goes unaccounted."""
    pursuit, bindings = two_docx
    union = set()
    for binding in bindings:
        facts = preview_docx_writeback(pursuit, at=AT, binding=binding)
        recorded = {c["slot_id"] for c in facts["cells"]}
        assert recorded, binding["file"]
        assert not (union & recorded)  # no slot claimed by two records
        union |= recorded
    container = pursuit.read_artifact("slots.json")
    expected = {s["slot_id"] for s in container["slots"]
                if not s.get("is_header")}
    assert union == expected


def test_bare_call_on_a_multi_docx_container_names_the_binding_door(
        two_docx):
    """B74§3g's refusal is DELETED with the capability landing; what
    remains is a caller-contract guard — a bare call cannot say which
    file it means, and the message points at declared_deliverables."""
    pursuit, _ = two_docx
    with pytest.raises(ContractError, match="OWN binding"):
        preview_docx_writeback(pursuit, at=AT)
