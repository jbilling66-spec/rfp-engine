"""Extraction contract: verbatim, loud, hidden-aware, date-honest."""

import datetime

import pytest
from openpyxl import Workbook

from engine.intake import UnreadableRfp, extract
from tests.fixtures.intake_twins import (
    HIDDEN_ROW_DIRECTIVE,
    HIDDEN_SHEET_DIRECTIVE,
    build_minimal_pdf,
)
from tests.fixtures.twins import _freeze
from tests.intake.test_fixtures import FIXTURES


# ------------------------------------------------------------------- xlsx


def test_sheet_labels_verbatim_including_trailing_space():
    doc = extract(FIXTURES / "structured-twin.xlsx")
    assert "## Sheet: 2. Integration \n" in doc.text + "\n"  # EC-1: space is real
    assert "## Sheet: 1. Company Background" in doc.text


def test_merged_range_annotated_and_formulas_preserved_as_text():
    doc = extract(FIXTURES / "structured-twin.xlsx")
    assert "Technical capability (30%) (merged D2:D4)" in doc.text  # EC-4
    assert "=SUM(B2:B4)" in doc.text  # EC-5b: formula TEXT, not a thinned value
    assert "='1. Company Background'!B2" in doc.text  # EC-5a


def test_hidden_sheet_and_row_extracted_and_marked():
    doc = extract(FIXTURES / "hidden-twin.xlsx")
    assert HIDDEN_SHEET_DIRECTIVE in doc.text
    assert HIDDEN_ROW_DIRECTIVE in doc.text
    assert "[hidden sheet]" in doc.text
    assert "[hidden row]" in doc.text
    hidden_text = " ".join(seg["text"] for seg in doc.hidden_segments)
    assert HIDDEN_SHEET_DIRECTIVE in hidden_text
    assert HIDDEN_ROW_DIRECTIVE in hidden_text
    locations = {seg["location"] for seg in doc.hidden_segments}
    assert "hidden-twin.xlsx: Internal Notes" in locations
    assert "hidden-twin.xlsx: Vendor Questions" in locations


def test_datetime_cells_render_iso_never_str(tmp_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Schedule"
    ws["A1"] = "Kickoff"
    ws["B1"] = datetime.datetime(2026, 3, 1)  # v1 rendered "2026-03-01 00:00:00"
    ws["B2"] = datetime.datetime(2026, 3, 1, 14, 30)
    path = tmp_path / "dates.xlsx"
    _freeze(wb, path)
    doc = extract(path)
    assert "2026-03-01 |" in doc.text
    assert "2026-03-01 00:00:00" not in doc.text
    assert "2026-03-01T14:30:00" in doc.text


def test_nofill_twin_extracts_nonempty_text():
    # EC-3 at P3 altitude: extraction is never silently empty on this file
    doc = extract(FIXTURES / "nofill-twin.xlsx")
    assert "Vendor shall describe its project management approach." in doc.text
    assert "Vendor shall describe its training and OCM approach." in doc.text


# -------------------------------------------------------------------- pdf


def test_pdf_page_markers_and_date_candidates():
    doc = extract(FIXTURES / "pdf-twin.pdf")
    assert "[page 1]" in doc.text
    assert "[page 2]" in doc.text
    by_text = {c["date_text"]: c for c in doc.date_candidates}
    assert by_text["August 8, 2026"]["date"] == "2026-08-08"
    assert by_text["August 29, 2026"]["date"] == "2026-08-29"
    assert by_text["August 29, 2026"]["location"] == "pdf-twin.pdf p1"


def test_pdf_empty_page_warns_all_empty_raises(tmp_path):
    partial = build_minimal_pdf(tmp_path / "partial.pdf", [["Some content here."], []])
    doc = extract(partial)
    assert any("page 2 produced no text" in w for w in doc.warnings)
    empty = build_minimal_pdf(tmp_path / "empty.pdf", [[], []])
    with pytest.raises(UnreadableRfp, match="no extractable text"):
        extract(empty)


def test_garbage_pdf_raises_unreadable(tmp_path):
    bad = tmp_path / "bad.pdf"
    bad.write_bytes(b"\x00\x01 this is not a pdf \xff")
    with pytest.raises(UnreadableRfp):
        extract(bad)


# ------------------------------------------------------------------- docx


def test_docx_headings_and_tables_in_document_order(tmp_path):
    from docx import Document

    document = Document()
    document.add_heading("Scope of Services", level=1)
    document.add_paragraph("The vendor shall provide implementation services.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Criterion"
    table.rows[0].cells[1].text = "Weight"
    document.add_paragraph("Submit questions in writing.")
    path = tmp_path / "sample.docx"
    document.save(path)
    doc = extract(path)
    lines = doc.text.splitlines()
    assert "# Scope of Services" in lines
    assert "| Criterion | Weight |" in lines
    # interleaving preserved: paragraph AFTER the table stays after it
    assert lines.index("| Criterion | Weight |") < lines.index("Submit questions in writing.")


# ------------------------------------------------------- text, refusals


def test_cp1252_text_file_decodes_with_warning(tmp_path):
    path = tmp_path / "notes.txt"
    path.write_bytes("Café budget — résumé attached".encode("cp1252"))
    doc = extract(path)
    assert "Café" in doc.text
    assert any("cp1252" in w for w in doc.warnings)


def test_missing_pptx_and_unknown_suffix_raise_unreadable(tmp_path):
    with pytest.raises(UnreadableRfp, match="not found"):
        extract(tmp_path / "ghost.pdf")
    pptx = tmp_path / "deck.pptx"
    pptx.write_bytes(b"anything")
    with pytest.raises(UnreadableRfp, match="B20"):
        extract(pptx)
    weird = tmp_path / "data.csv"
    weird.write_text("a,b,c")
    with pytest.raises(UnreadableRfp, match="unsupported format"):
        extract(weird)
