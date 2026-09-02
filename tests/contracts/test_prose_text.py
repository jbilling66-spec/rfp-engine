"""engine/contracts/text.py — the typed control-character refusal at the
prose boundary (P26a, P2-29b). The exit doors already close a run on
python-docx's ValueError (P25 item 8); this is the entry-side check every
prose writer asks first."""

import pytest
from docx import Document

from engine.contracts import check_prose


def test_clean_prose_and_the_xml_whitespace_trio_pass():
    assert check_prose("Plain prose.\nSecond line.\tTabbed.\r\n") is None
    assert check_prose("Unicode is fine: café — “quotes” ✓") is None


@pytest.mark.parametrize("bad", ["bad\x0bchar", "form\x0cfeed", "nul\x00",
                                 "esc\x1b[0m", "del\x7f"])
def test_control_characters_are_named_by_codepoint_never_reproduced(bad):
    reason = check_prose(bad)
    assert reason and reason.startswith("control character U+")
    assert all(ch not in reason for ch in bad if ord(ch) < 32 or ord(ch) == 127)


def test_the_rule_matches_what_python_docx_refuses():
    """The predicate is calibrated against the library it protects: every
    string it passes, python-docx accepts; the one it names, python-docx
    raises on."""
    doc = Document()
    doc.add_paragraph("tab\tok\nline")
    with pytest.raises(ValueError):
        doc.add_paragraph("bad\x0bchar")
    assert check_prose("bad\x0bchar") is not None


def test_non_text_is_refused_typed():
    assert "expected text" in check_prose(42)
