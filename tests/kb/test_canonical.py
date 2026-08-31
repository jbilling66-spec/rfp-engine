"""C3/C4 (P13): the L1 canonical model — round-trip fidelity, content-
addressed ids, byte-determinism across directories (the model file joins
the byte-golden seed store at C8, so any nondeterminism here flakes that
golden), refusal of shape drift, and the same-shape twin proof: three
input routes (runtime-built DOCX, markdown, scripted ExtractionView)
resolve to ONE element shape (WP13 R1).

Evidence limit, stated at B59: the ExtractionView route is proven over a
SCRIPTED view here — the real docling conversion evidence rides the A1
§A2 buyer-corpus rerun. This is the venv proof, not a docling proof."""

from pathlib import Path

import pytest

from engine.contracts import ContractError
from engine.extraction.model import ExtractionView, FigureView
from engine.kb.canonical import (CanonicalDoc, Chunk, Element, doc_id_for,
                                 elements_from_extraction_view,
                                 elements_from_markdown, read_model,
                                 source_hash_for, view_extraction_status,
                                 write_model)
from engine.kb.read import read_source


def _model(source: bytes = b"synthetic northwind response") -> CanonicalDoc:
    return CanonicalDoc(
        doc_id=doc_id_for(source),
        source_hash=source_hash_for(source),
        extractor="python-docx",
        extraction_fingerprint="ext_0123456789abcdef",
        extraction_status="clean",
        media={"images": 1},
        elements=[
            Element(kind="heading", text="6.0 Accelerators", level=1),
            Element(kind="paragraph", text="The migration factory ran."),
            Element(kind="table_row", text="Criterion | Weight"),
            Element(kind="figure", text="", figure_class="chart"),
        ],
        chunks=[
            Chunk(doc_path=["6.0 Accelerators"], elements=(1, 4),
                  chars=42, pages=[12], kb_id="kb_0123456789"),
            Chunk(doc_path=["6.0 Accelerators"], elements=(3, 4), chars=0),
        ],
    )


def test_round_trip_is_lossless(tmp_path):
    model = _model()
    write_model(tmp_path, model)
    assert read_model(tmp_path, model.doc_id) == model


def test_bytes_deterministic_across_directories(tmp_path):
    model = _model()
    a = write_model(tmp_path / "one", model).read_bytes()
    b = write_model(tmp_path / "two", model).read_bytes()
    assert a == b


def test_doc_id_is_content_addressed():
    assert doc_id_for(b"same") == doc_id_for(b"same")
    assert doc_id_for(b"same") != doc_id_for(b"different")
    assert doc_id_for(b"same").startswith("cd_")
    assert len(doc_id_for(b"same")) == 15  # cd_ + 12 hex


def test_write_refuses_shape_drift(tmp_path):
    model = _model()
    model.elements[0].kind = "footnote"
    with pytest.raises(ContractError):
        write_model(tmp_path, model)
    assert not (tmp_path / "canonical").exists(), \
        "a refused model must leave nothing behind"


def test_from_dict_refuses_unknown_keys():
    payload = _model().to_dict()
    payload["parsed_at"] = "2026-08-23"  # clocks stay out, loudly
    with pytest.raises(ValueError):
        CanonicalDoc.from_dict(payload)


def test_optional_chunk_fields_omitted_when_empty(tmp_path):
    """Sparse serialization is part of the byte contract: an empty pages
    list or an unminted kb_id must not appear as noise in the file."""
    model = _model()
    write_model(tmp_path, model)
    raw = (Path(tmp_path) / "canonical"
           / f"{model.doc_id}.json").read_text(encoding="utf-8")
    payload_chunks = read_model(tmp_path, model.doc_id).chunks
    assert payload_chunks[1].pages == [] and payload_chunks[1].kb_id is None
    assert raw.count('"kb_id"') == 1


# ------------------------------------------- C4: two producers, one model

_TWIN_MARKDOWN = """\
# 6.0 Accelerators

The migration factory ran on schedule.
It processed nine waves without a rollback.

## 6.1 Tooling

| Criterion | Weight |
| --- | --- |
| Technical approach | 40% |

The toolchain was operated by the firm's own staff.
"""


def _twin_docx(tmp_path) -> Path:
    import docx

    document = docx.Document()
    document.add_heading("6.0 Accelerators", level=1)
    document.add_paragraph("The migration factory ran on schedule.\n"
                           "It processed nine waves without a rollback.")
    document.add_heading("6.1 Tooling", level=2)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Criterion"
    table.rows[0].cells[1].text = "Weight"
    table.rows[1].cells[0].text = "Technical approach"
    table.rows[1].cells[1].text = "40%"
    document.add_paragraph("The toolchain was operated by the firm's "
                           "own staff.")
    path = tmp_path / "twin.docx"
    document.save(str(path))
    return path


def _twin_view() -> ExtractionView:
    """The PDF shape: same content split across two pages, arriving as a
    scripted view through the P12 seam."""
    page_one, page_two = _TWIN_MARKDOWN.split("## 6.1 Tooling")
    return ExtractionView(
        text=_TWIN_MARKDOWN, pages=2,
        page_texts=[page_one, "## 6.1 Tooling" + page_two])


def _shape(elements) -> list[tuple]:
    return [(e.kind, e.text, e.level) for e in elements]


def test_three_routes_one_shape(tmp_path):
    """WP13 R1: DOCX, markdown, and the extraction-view route resolve to
    the same element shape — the model is stack-agnostic."""
    from_markdown = elements_from_markdown(_TWIN_MARKDOWN)
    from_docx = read_source(_twin_docx(tmp_path)).elements
    from_view = elements_from_extraction_view(_twin_view())
    assert _shape(from_markdown) == _shape(from_docx) == _shape(from_view)
    kinds = [e.kind for e in from_markdown]
    assert kinds == ["heading", "paragraph", "heading",
                     "table_row", "table_row", "paragraph"]


def test_view_route_keeps_page_provenance():
    pages = [e.page for e in elements_from_extraction_view(_twin_view())]
    assert pages == [1, 1, 2, 2, 2, 2]


def test_view_figures_become_figure_elements():
    view = _twin_view()
    view.figures = [
        FigureView(classes=[{"label": "chart", "confidence": 0.9},
                            {"label": "photo", "confidence": 0.2}]),
        FigureView(classes=[{"label": "bar_chart_exotic",
                             "confidence": 0.7}]),
        FigureView(classes=[]),
    ]
    figures = [e for e in elements_from_extraction_view(view)
               if e.kind == "figure"]
    assert [f.figure_class for f in figures] == ["chart", "other", None]


def test_partial_success_maps_to_degraded():
    view = _twin_view()
    assert view_extraction_status(view) == "clean"
    view.status = "partial_success"
    assert view_extraction_status(view) == "degraded"


def test_html_comment_lines_stay_out_of_the_model():
    """The eval harness's doc_meta comment names a client — harness
    metadata must never become an L1 element."""
    text = "<!-- client: Northwind | outcome: won -->\n" + _TWIN_MARKDOWN
    assert _shape(elements_from_markdown(text)) == \
        _shape(elements_from_markdown(_TWIN_MARKDOWN))


def test_markdown_reader_route_carries_elements(tmp_path):
    path = tmp_path / "twin.md"
    path.write_text(_TWIN_MARKDOWN, encoding="utf-8")
    source = read_source(path)
    assert source.text == _TWIN_MARKDOWN  # flat text stays byte-unchanged
    assert _shape(source.elements) == \
        _shape(elements_from_markdown(_TWIN_MARKDOWN))
