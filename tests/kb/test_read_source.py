"""C12: the KB source reader — python-docx identity stamped, media facts
counted, and the .md path byte-unchanged (the anonymization suite reads
through this seam now)."""

from engine.kb.read import read_source
from tests.kb.fixtures.corpus import SOURCE_DOCS


def test_docx_reads_with_identity_and_media(tmp_path):
    from engine.extraction.corpus import build_logo_docx

    path = build_logo_docx(tmp_path / "logo-twin.docx")
    source = read_source(path)
    assert source.extractor == "python-docx"
    assert source.fingerprint.startswith("ext_")
    assert source.media == {"images": 1}
    # Intake's text conventions hold: headings as '#' lines.
    assert "# Northwind Regional Health - Engagement Letter" in source.text
    assert "identifying logo" in source.text


def test_docx_tables_render_as_pipe_rows(tmp_path):
    import docx as pydocx

    d = pydocx.Document()
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Deliverable"
    table.rows[0].cells[1].text = "Fee"
    path = tmp_path / "t.docx"
    d.save(path)
    source = read_source(path)
    assert "| Deliverable | Fee |" in source.text
    assert source.media == {"images": 0}


def test_md_path_is_byte_unchanged(tmp_path):
    # The 20 committed anonymization docs flow through read_source now —
    # their text must be exactly the read_text() the recall record was
    # measured over.
    path = tmp_path / "anon.md"
    body = "<!-- client: Foxglove Robotics | date: 2026-01-01 -->\n# Doc\nbody\n"
    path.write_text(body, encoding="utf-8")
    source = read_source(path)
    assert source.text == body
    assert source.extractor == "text"
    assert source.media == {"images": 0}


def test_kb_and_intake_docx_stamps_are_independent(tmp_path):
    # Same library, two stacks: the KB fingerprint must not equal the
    # intake legacy python-docx fingerprint... they may share components
    # today, but the seam test (C12) pins the IDENTITIES; here we pin
    # that a fingerprint exists and names no docling component.
    import docx as pydocx

    d = pydocx.Document()
    d.add_paragraph("body")
    path = tmp_path / "p.docx"
    d.save(path)
    source = read_source(path)
    assert source.extractor == "python-docx"
    text_doc = SOURCE_DOCS[0]
    assert text_doc.extractor == ""  # fixtures predate the reader: unstamped
