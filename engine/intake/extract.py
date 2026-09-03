"""Deterministic buyer-document text extraction (P3; v1 engine/ingest is the
oracle, reimplemented not ported).

Loud-failure contract: UnreadableRfp on a missing file, unsupported format,
unparseable bytes, or empty extraction — never a silently empty brief. Three
v1 holes fixed by design: cp1252 fallback before failing on text files, an
explicit warning for PDF pages that yield no text (a half-scanned document
must not silently thin), and datetime cells rendered via isoformat (never
str(cell) → "2026-03-01 00:00:00").

Hidden workbook SHEETS, ROWS and COLUMNS are extracted AND marked (v1
iterated hidden sheets and rows silently — exactly the surface an
injection screen must see; columns joined at P26b-1, P1-26).
"""

import contextlib
import datetime
import importlib.metadata
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

import pypdf
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from engine.extraction.backend import ExtractionFailed
from engine.extraction.fingerprint import (
    extraction_fingerprint,
    manifest_digest,
    stack_fingerprint,
)
from engine.extraction.media import media_findings
from engine.structure.zipguard import check_office_zip


class UnreadableRfp(Exception):
    def __init__(self, path, why: str):
        self.path = str(path)
        self.why = why
        super().__init__(f"{path}: {why}")


@dataclass
class ExtractedDoc:
    file: str
    format: str  # pdf | docx | xlsx | other  (intake.documents enum)
    text: str  # full extracted text, hidden segments included and marked
    warnings: list[str] = field(default_factory=list)
    hidden_segments: list[dict] = field(default_factory=list)  # {text, location}
    date_candidates: list[dict] = field(default_factory=list)  # {date_text, date, location}
    # C9 (B57): every document says which stack read it. Flags are short
    # slugs (reasons go to warnings); degraded means a weaker extraction
    # than the adopted primary for this format — it still ingests (C10).
    extractor: str = ""  # docling | pypdf | python-docx | openpyxl | text
    extraction_fingerprint: str = ""
    extraction_degraded: bool = False
    extraction_flags: list[str] = field(default_factory=list)
    grids: list | None = None  # docling table views [{grid, merges}] (C10 diff)
    sidecar: dict | None = None  # DOCX fills/comments the layer recovers (B57)


def location_of(text: str, pos: int, file: str) -> str:
    """Nearest preceding [page N] / '## Sheet:' marker — the single source
    for source_location strings (screen and date scan both use it)."""
    marker = None
    for match in re.finditer(r"^\[page (\d+)\]$|^## Sheet: (.*)$", text[:pos], flags=re.M):
        marker = match
    if marker is None:
        return file
    if marker.group(1):
        return f"{file} p{marker.group(1)}"
    return f"{file}: {marker.group(2)}"


# ------------------------------------------------------------------ formats


def _cell_text(value) -> str:
    if isinstance(value, datetime.datetime):
        if value.time() == datetime.time(0, 0):
            return value.date().isoformat()
        return value.isoformat()
    if isinstance(value, datetime.date):
        return value.isoformat()
    return str(value).replace("\n", " ")


def _extract_xlsx(path: Path, doc: ExtractedDoc) -> None:
    try:
        check_office_zip(path)  # P0-8: the container before the parser
        wb = load_workbook(path, data_only=False)  # hidden state + formula TEXT
    except Exception as exc:  # openpyxl raises a zoo of types on bad bytes
        raise UnreadableRfp(path, f"unparseable xlsx: {exc}") from exc
    lines: list[str] = []
    cached_wb = None  # P1-25: loaded once, only if a formula cell exists
    for ws in wb.worksheets:
        lines.append(f"## Sheet: {ws.title}")  # verbatim — trailing spaces are real
        sheet_hidden = ws.sheet_state != "visible"
        if sheet_hidden:
            lines.append("[hidden sheet]")
        merged_anchors = {
            (rng.min_row, rng.min_col): str(rng) for rng in ws.merged_cells.ranges
        }
        # P1-26 (P26b-1, B112): a hidden column dimension covers min..max.
        hidden_cols: set[int] = set()
        for dim in ws.column_dimensions.values():
            if dim.hidden:
                hidden_cols.update(range(dim.min, dim.max + 1))
        col_segments: dict[int, list[str]] = {}
        for row in ws.iter_rows():
            texts = []
            for cell in row:
                if cell.value is None:
                    continue
                rendered = _cell_text(cell.value)
                if cell.column in hidden_cols:
                    col_segments.setdefault(cell.column, []).append(rendered)
                    rendered = f"[hidden col] {rendered}"
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    # P1-25 (P26b-1, B112): a formula is MARKED so the
                    # model never mistakes its source for what the
                    # human sees; the cached value (what Excel saved)
                    # rides beside it when one exists, and its absence
                    # is a warning. The source text stays (EC-5: never
                    # thinned; the KB lane refuses, intake reads).
                    if cached_wb is None:
                        cached_wb = load_workbook(path, data_only=True)
                    cached = cached_wb[ws.title][cell.coordinate].value
                    if cached is None or str(cached) == "":
                        rendered = f"[formula, no cached value] {rendered}"
                        doc.warnings.append(
                            f"{ws.title}!{cell.coordinate}: formula cell without "
                            "a cached value — the human's view of this cell is "
                            "unknown to the engine (paste values, or open and "
                            "save in Excel)")
                    else:
                        rendered = f"[formula {rendered} → {_cell_text(cached)}]"
                anchor = merged_anchors.get((cell.row, cell.column))
                if anchor:
                    rendered += f" (merged {anchor})"
                texts.append(rendered)
            if not texts:
                continue
            row_hidden = ws.row_dimensions[row[0].row].hidden
            line = "| " + " | ".join(texts) + " |"
            if sheet_hidden or row_hidden:
                doc.hidden_segments.append(
                    {"text": " ".join(texts), "location": f"{doc.file}: {ws.title}"}
                )
            if row_hidden:
                line = "[hidden row] " + line
            lines.append(line)
        for col, col_texts in sorted(col_segments.items()):
            doc.hidden_segments.append({
                "text": " ".join(col_texts),
                "location": f"{doc.file}: {ws.title}!{get_column_letter(col)}",
            })
    doc.text = "\n".join(lines)


class _PypdfWarnings(logging.Handler):
    """P2-26 (P26b-1, B112): pypdf reports recovery (a bad startxref, a
    rebuilt xref, an odd object) on its `pypdf` loggers and nothing in
    this engine listened — a partially-recovered PDF looked clean. This
    handler is attached for the duration of ONE read and removed in
    `finally`; the engine's first and only use of `logging`, scoped, with
    no global configuration."""

    def __init__(self) -> None:
        super().__init__(level=logging.WARNING)
        self.messages: list[str] = []

    def emit(self, record: logging.LogRecord) -> None:
        self.messages.append(record.getMessage())


@contextlib.contextmanager
def _capture_pypdf_warnings():
    logger = logging.getLogger("pypdf")
    handler = _PypdfWarnings()
    prior_level = logger.level
    logger.addHandler(handler)
    if logger.getEffectiveLevel() > logging.WARNING:
        logger.setLevel(logging.WARNING)
    try:
        yield handler.messages
    finally:
        logger.removeHandler(handler)
        logger.setLevel(prior_level)


def _extract_pdf(path: Path, doc: ExtractedDoc) -> None:
    try:
        with _capture_pypdf_warnings() as recovered:
            reader = pypdf.PdfReader(path)
            pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:
        raise UnreadableRfp(path, f"unparseable pdf: {exc}") from exc
    for message in recovered:
        doc.warnings.append(f"pdf recovery: {message}")
    for number, text in enumerate(pages, start=1):
        if not text.strip():
            doc.warnings.append(f"page {number} produced no text (image-only or empty)")
    doc.text = "\n".join(
        f"[page {number}]\n{text}" for number, text in enumerate(pages, start=1)
    )


def _extract_docx(path: Path, doc: ExtractedDoc) -> None:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        check_office_zip(path)  # P0-8: the container before the parser
        document = Document(path)
    except Exception as exc:
        raise UnreadableRfp(path, f"unparseable docx: {exc}") from exc
    from engine.structure.docx_parts import header_footer_text, text_box_text

    # P2-27 (P26b-1, B112): headers first, footers last, text boxes
    # inline after their anchor — every part the injection screen must
    # see, marked so the model knows where the words sat.
    parts = header_footer_text(document)
    lines: list[str] = [f"[header] {text}" for kind, text in parts
                        if kind == "header"]
    # iter_inner_content preserves paragraph/table interleaving —
    # .paragraphs/.tables lose document order (v1 lesson)
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            style = item.style.name if item.style is not None else ""
            match = re.fullmatch(r"Heading (\d)", style or "")
            if match and item.text.strip():
                lines.append("#" * int(match.group(1)) + " " + item.text)
            elif item.text.strip():
                lines.append(item.text)
            lines.extend(f"[text box] {text}" for text in text_box_text(item._p))
        elif isinstance(item, Table):
            for row in item.rows:
                texts = [c.text.replace("\n", " ") for c in row.cells if c.text.strip()]
                if texts:
                    lines.append("| " + " | ".join(texts) + " |")
                for cell in row.cells:
                    lines.extend(f"[text box] {text}"
                                 for text in text_box_text(cell._tc))
    lines.extend(f"[footer] {text}" for kind, text in parts if kind == "footer")
    doc.text = "\n".join(lines)


def _extract_text_file(path: Path, doc: ExtractedDoc) -> None:
    raw = path.read_bytes()
    for encoding in ("utf-8", "cp1252"):
        try:
            doc.text = raw.decode(encoding)
            if encoding != "utf-8":
                doc.warnings.append(f"decoded as {encoding} (not valid utf-8)")
            return
        except UnicodeDecodeError:
            continue
    raise UnreadableRfp(path, "undecodable text file (not utf-8 or cp1252)")


# --------------------------------------------------------------- date scan

_MONTHS = (
    "January|February|March|April|May|June|July|August|September|October|November|December"
)
_DATE_PATTERNS = [
    (rf"\b(?:{_MONTHS}) \d{{1,2}}, \d{{4}}\b", "%B %d, %Y"),
    (r"\b\d{4}-\d{2}-\d{2}\b", "%Y-%m-%d"),
    (r"\b\d{1,2}/\d{1,2}/\d{4}\b", "%m/%d/%Y"),
]


def parse_date(text: str) -> str | None:
    """First parseable date in `text` as ISO, else None. Same pattern table
    as the document scan — one implementation of the date rule."""
    best: tuple[int, str, str] | None = None
    for pattern, fmt in _DATE_PATTERNS:
        match = re.search(pattern, text)
        if match and (best is None or match.start() < best[0]):
            best = (match.start(), match.group(0), fmt)
    if best is None:
        return None
    try:
        return datetime.datetime.strptime(best[1], best[2]).date().isoformat()
    except ValueError:
        return None


def _scan_dates(doc: ExtractedDoc) -> None:
    for pattern, fmt in _DATE_PATTERNS:
        for match in re.finditer(pattern, doc.text):
            date_text = match.group(0)
            try:
                parsed = datetime.datetime.strptime(date_text, fmt).date().isoformat()
            except ValueError:
                parsed = None
                doc.warnings.append(f"date-like text did not parse: {date_text!r}")
            doc.date_candidates.append(
                {
                    "date_text": date_text,
                    "date": parsed,
                    "location": location_of(doc.text, match.start(), doc.file),
                }
            )


# --------------------------------------------------------------- dispatcher

_FORMATS = {
    ".pdf": ("pdf", _extract_pdf),
    ".docx": ("docx", _extract_docx),
    ".xlsx": ("xlsx", _extract_xlsx),
    ".md": ("other", _extract_text_file),
    ".txt": ("other", _extract_text_file),
}

# Legacy stack identities (the C12 seam: every stamp names a real stack).
_LEGACY_IDENTITY = {
    ".pdf": ("pypdf", "pypdf"),
    ".docx": ("python-docx", "python-docx"),
    ".xlsx": ("openpyxl", "openpyxl"),
    ".md": ("text", None),
    ".txt": ("text", None),
}


def _stamp_legacy(doc: ExtractedDoc, suffix: str) -> None:
    name, dist = _LEGACY_IDENTITY[suffix]
    version = importlib.metadata.version(dist) if dist else "stdlib"
    doc.extractor = name
    doc.extraction_fingerprint = stack_fingerprint(
        name, {"extractor_version": version}
    )


def _extract_docling(path: Path, fmt: str, backend) -> ExtractedDoc:
    """The adopted primary for pdf/docx intake (B57). PDFs rebuild the
    [page N]-marked text from per-page exports so location_of() and the
    deadline source_location search behave exactly as on the legacy path;
    DOCX takes the markdown export (headings arrive as '#' lines)."""
    view = backend.convert(path)  # ExtractionFailed propagates to extract()
    doc = ExtractedDoc(file=path.name, format=fmt, text="")
    if fmt == "pdf":
        pages = view.page_texts or []
        for number, text in enumerate(pages, start=1):
            if not text.strip():
                doc.warnings.append(
                    f"page {number} produced no text (image-only or empty)"
                )
        doc.text = "\n".join(
            f"[page {number}]\n{text}" for number, text in enumerate(pages, start=1)
        )
    else:
        doc.text = view.text
    doc.extractor = "docling"
    doc.extraction_fingerprint = extraction_fingerprint(
        view.docling_version, manifest_digest()
    )
    doc.grids = [{"grid": t.grid, "merges": t.merges} for t in view.grids]
    doc.sidecar = view.sidecar
    if view.status == "partial_success":
        # The worker surfaces what the gate used to swallow; C10 carries
        # it into the run log and the ingest report.
        doc.extraction_degraded = True
        doc.extraction_flags.append("partial_extraction")
    # C11: classified identity-bearing figures (logo/signature) flag the
    # document — computed since C5, consumed here.
    doc.extraction_flags.extend(media_findings(view.figures))
    # C13 (B57 accept-with-flag): multicolumn layout -> mandatory review,
    # never degraded — the read is fine, the reading ORDER is the risk.
    if view.multicolumn_pages:
        doc.extraction_flags.append("multicolumn_layout")
    return doc


def extract(path: Path, *, backend=None) -> ExtractedDoc:
    path = Path(path)
    if not path.is_file():
        raise UnreadableRfp(path, "file not found")
    suffix = path.suffix.lower()
    if suffix == ".pptx":
        raise UnreadableRfp(path, "pptx deferred pending dependency review — B20")
    if suffix not in _FORMATS:
        raise UnreadableRfp(path, f"unsupported format {suffix!r}")
    fmt, extractor = _FORMATS[suffix]
    if backend is not None and suffix in (".pdf", ".docx"):
        try:
            doc = _extract_docling(path, fmt, backend)
        except ExtractionFailed as exc:
            # Document-level failure only: the legacy stack reads it,
            # stamped degraded — it still ingests, flagged (C10). An
            # unavailable ENVIRONMENT never reaches here (backend
            # construction refuses, the owner's call — B58).
            doc = ExtractedDoc(file=path.name, format=fmt, text="")
            extractor(path, doc)
            _stamp_legacy(doc, suffix)
            doc.extraction_degraded = True
            doc.extraction_flags.append("docling_fallback")
            doc.warnings.append(
                f"docling extraction failed; legacy fallback: {str(exc)[:200]}"
            )
    else:
        doc = ExtractedDoc(file=path.name, format=fmt, text="")
        extractor(path, doc)
        _stamp_legacy(doc, suffix)
        if backend is None and suffix in (".pdf", ".docx"):
            # No backend wired at this call site (tests, legacy callers):
            # the doc is honest about being below the adopted primary.
            doc.extraction_degraded = True
            doc.extraction_flags.append("legacy_extractor")
    content = re.sub(
        r"^\[page \d+\]$|^## Sheet: .*$|^\[hidden sheet\]$", "", doc.text, flags=re.M
    )
    if not re.search(r"\w", content):
        raise UnreadableRfp(path, "no extractable text")
    _scan_dates(doc)
    return doc
