"""Production conversion worker (C9) — the same docling path the §A2 gate
proved, stripped of gate-only evidence machinery.

Three deliberate differences from gate.convert_worker:
- ONE conversion per document. The gate's warm second pass exists to score
  p95 honestly; production artifacts carry no timing at all (byte-identity
  across kill/resume, brief.py).
- `partial_success` is RETURNED as status, never swallowed: the intake
  adapter turns it into the degraded flag (C10). The gate accepts partial
  because its scorers judge the output; production must say so out loud.
- Per-page text (`page_texts`) is exported for PDFs so intake keeps its
  `[page N]` markers — location_of() and the deadline source_location
  search work unchanged on the docling path.

The sandbox child target is `convert_production`; `run_production_conversion`
wraps it in the C2 jail. Callers verify weights BEFORE constructing/invoking
(the B51 construction-refusal — backend.py owns it); this module only jails
and converts. The `__main__` entry is the DockerBackend's in-container
transport: verify → jailed convert → one JSON object on stdout.
"""

from __future__ import annotations

from pathlib import Path

PRODUCTION_TIMEOUT_S = 900.0
PRODUCTION_MEM_MB = 12288  # matches the gate's jail ceiling


def assemble_view(doc, source: Path) -> dict:
    """Build the serializable view halves shared by the gate worker and
    production: grids, headings, figures, comment texts, and the DOCX
    sidecar. Pure over the docling document + source path; the caller adds
    text/pages/status and (gate only) timing."""
    grids = []
    for table in doc.tables:
        data = table.data
        grid = [["" for _ in range(data.num_cols)] for _ in range(data.num_rows)]
        merges = []
        for cell in data.table_cells:
            for r in range(cell.start_row_offset_idx, cell.end_row_offset_idx):
                for c in range(cell.start_col_offset_idx, cell.end_col_offset_idx):
                    grid[r][c] = cell.text
            if (cell.end_row_offset_idx - cell.start_row_offset_idx) > 1 or (
                cell.end_col_offset_idx - cell.start_col_offset_idx
            ) > 1:
                merges.append(
                    [
                        [cell.start_row_offset_idx, cell.start_col_offset_idx],
                        [cell.end_row_offset_idx - 1, cell.end_col_offset_idx - 1],
                    ]
                )
        grids.append({"grid": grid, "merges": merges})

    headings = [
        [getattr(item, "level", 1), item.text]
        for item in getattr(doc, "texts", [])
        if getattr(getattr(item, "label", None), "value", "") == "section_header"
    ]
    # docling carries DOCX comment TEXT as floating text items (probed
    # 2026-08-15: "[author: ...]: <text>"); the anchor cell is lost.
    native_comment_texts = [
        str(item.text)
        for item in getattr(doc, "texts", [])
        if str(getattr(item, "text", "")).startswith("[author:")
    ]

    # The DOCX sidecar channel (§A2.4 subordinate call, adopted at B57):
    # docling's model carries NO cell shading — the extraction LAYER
    # recovers w:shd and comments via python-docx beside the docling parse.
    sidecar: dict = {"fills": {}, "comment_texts": []}
    if source.suffix.lower() == ".docx":
        import re as _re
        import zipfile as _zip

        import docx as pydocx
        from docx.oxml.ns import qn

        pdoc = pydocx.Document(str(source))
        for t_idx, t in enumerate(pdoc.tables):
            per: dict = {}
            for r, row in enumerate(t.rows):
                for c, cell in enumerate(row.cells):
                    for shd in cell._tc.iter(qn("w:shd")):
                        fill = shd.get(qn("w:fill"))
                        if fill and fill.lower() != "auto" and [r, c] not in per.setdefault(fill, []):
                            per[fill].append([r, c])
            if per:
                sidecar["fills"][str(t_idx)] = per
        with _zip.ZipFile(source) as zf:
            if "word/comments.xml" in zf.namelist():
                xml = zf.read("word/comments.xml").decode("utf-8")
                sidecar["comment_texts"] = [
                    "".join(_re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml))
                ]

    figures = [
        {
            "classes": [
                {"label": ann.predicted_classes[0].class_name,
                 "confidence": ann.predicted_classes[0].confidence}
                for ann in getattr(pic, "annotations", [])
                if getattr(ann, "predicted_classes", None)
            ]
        }
        for pic in getattr(doc, "pictures", [])
    ]

    return {
        "grids": grids,
        "headings": headings,
        "figures": figures,
        "native_comment_texts": native_comment_texts,
        "sidecar": sidecar,
    }


def _page_texts(doc, grids: list) -> list[str]:
    """Per-page plain text from provenance: text items plus pipe-row table
    renderings (the legacy extractor's table convention), grouped by the
    1-based page each item sits on. Deterministic — document order within
    each page."""
    pages = len(getattr(doc, "pages", {})) or 1
    per: dict[int, list[str]] = {n: [] for n in range(1, pages + 1)}
    for item in getattr(doc, "texts", []):
        prov = getattr(item, "prov", None)
        if prov:
            per.setdefault(prov[0].page_no, []).append(str(item.text))
    for t_idx, table in enumerate(getattr(doc, "tables", [])):
        prov = getattr(table, "prov", None)
        if prov and t_idx < len(grids):
            rows = grids[t_idx]["grid"]
            per.setdefault(prov[0].page_no, []).append(
                "\n".join("| " + " | ".join(row) + " |" for row in rows)
            )
    return ["\n".join(per.get(n, [])) for n in sorted(per)]


def convert_production(payload: dict) -> dict:
    """Sandbox child target: convert ONE document, return the typed-view
    dict (model.ExtractionView shape — no timing fields, ever)."""
    import docling as _docling
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption

    from engine.extraction.fingerprint import PIPELINE_OPTIONS

    source = Path(payload["file"])
    artifacts = payload["models_root"]

    if payload["mode"] == "vlm":
        from docling.datamodel.pipeline_options import VlmPipelineOptions
        from docling.pipeline.vlm_pipeline import VlmPipeline

        pdf_option = PdfFormatOption(
            pipeline_cls=VlmPipeline,
            pipeline_options=VlmPipelineOptions(artifacts_path=artifacts),
        )
    else:
        # The options here and the fingerprint's PIPELINE_OPTIONS are one
        # constant — the stamp can never disagree with the conversion.
        pdf_option = PdfFormatOption(
            pipeline_options=PdfPipelineOptions(
                artifacts_path=artifacts,
                do_ocr=PIPELINE_OPTIONS["do_ocr"],
                do_table_structure=PIPELINE_OPTIONS["do_table_structure"],
                do_picture_classification=PIPELINE_OPTIONS["do_picture_classification"],
            )
        )

    converter = DocumentConverter(
        format_options={InputFormat.PDF: pdf_option, InputFormat.IMAGE: pdf_option}
    )
    result = converter.convert(str(source), raises_on_error=False)
    status = str(getattr(result.status, "value", result.status))
    if status not in ("success", "partial_success"):
        details = "; ".join(
            str(getattr(e, "error_message", e)) for e in (result.errors or [])
        )
        raise RuntimeError(
            f"conversion {result.status}: {details[:600] or 'no detail recorded'}"
        )
    doc = result.document

    view = assemble_view(doc, source)
    pages = len(getattr(doc, "pages", {})) or 1
    is_pdf = source.suffix.lower() == ".pdf"
    view.update(
        {
            "text": doc.export_to_markdown(),
            "pages": pages,
            "page_texts": _page_texts(doc, view["grids"]) if is_pdf else None,
            "multicolumn_pages": _multicolumn_pages(doc, view["grids"])
            if is_pdf else [],
            "status": status,
            "docling_version": _docling.__version__,
        }
    )
    return view


def _multicolumn_pages(doc, grids: list) -> list[int]:
    """Both C13 signals (pure logic in engine/extraction/layout.py):
    text-item bbox clustering, plus tables the layout model built out of
    prose columns — the rendering the corpus probe actually gets."""
    from engine.extraction.layout import looks_like_prose_columns, multicolumn_pages

    page_boxes: dict[int, list] = {}
    for item in getattr(doc, "texts", []):
        prov = getattr(item, "prov", None)
        if prov and getattr(prov[0], "bbox", None) is not None:
            bb = prov[0].bbox
            page_boxes.setdefault(prov[0].page_no, []).append(
                [bb.l, bb.t, bb.r, bb.b]
            )
    flagged = set(multicolumn_pages(page_boxes))
    for t_idx, table in enumerate(getattr(doc, "tables", [])):
        prov = getattr(table, "prov", None)
        if (prov and t_idx < len(grids)
                and looks_like_prose_columns(grids[t_idx]["grid"])):
            flagged.add(prov[0].page_no)
    return sorted(flagged)


def run_production_conversion(path: Path, mode: str, workdir: Path,
                              timeout_s: float = PRODUCTION_TIMEOUT_S):
    """Jail convert_production for one document. Weights verification is
    the CALLER's obligation before this runs (backend construction / the
    __main__ transport below) — the jail governs execution, not supply."""
    from engine.extraction.sandbox import run_sandboxed
    from engine.extraction.weights import models_root

    return run_sandboxed(
        "engine.extraction.worker:convert_production",
        {"file": str(path), "mode": mode, "models_root": str(models_root() / "docling")},
        timeout_s=timeout_s,
        mem_mb=PRODUCTION_MEM_MB,
        workdir=workdir,
    )


if __name__ == "__main__":
    # DockerBackend transport (runs INSIDE the gate image): verify the
    # mounted weights against the committed manifest, convert in the jail,
    # emit exactly one JSON object on stdout. Argv: <file> <mode>.
    import json
    import sys
    import tempfile

    from engine.extraction.weights import verify_artifacts

    doc_path, doc_mode = Path(sys.argv[1]), sys.argv[2]
    verify_artifacts()
    with tempfile.TemporaryDirectory(prefix="extract-") as td:
        res = run_production_conversion(doc_path, doc_mode, Path(td))
    print(json.dumps(
        {"status": res.status, "view": res.result, "error": res.error}
    ))
    sys.exit(0 if res.status == "ok" else 1)
