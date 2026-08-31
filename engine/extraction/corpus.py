"""Synthetic §A2 gate corpus — generated at runtime, never committed (C4, B51).

The five A2.1 stand-ins (B44: the buyer-document half is A1's; these are
the synthetic half) plus the failure-behavior documents and the
logo/signature media docs for C11. Generate-at-runtime is the B40/D21
house pattern: the adversarial trait stays a line of Python and the
tripwire's extraction sweep gains no new binaries.

Every string here is synthetic by construction — the Northwind Regional
Health vocabulary the committed twins established. The committed answer
key is evals/extraction-gate/ground_truth.json, duplicated from the
grids below ON PURPOSE: the trait tests read each built document back
with the raw library and compare against the key, so a builder that
drifts from the key reddens instead of re-deriving its own truth.

Tables are built merge-first, then every grid position is written with
its expected text — continuation positions of a merged cell are the
same underlying cell, so duplicate writes are idempotent and the grid
constants read exactly like python-docx reads them back.

The PDF assembler mirrors tests/fixtures/intake_twins.py's hand-rolled
minimal PDF 1.4 (spec rule 3: a PDF-writing dependency would need human
review; this needs none). Duplicated rather than imported: engine code
does not import tests, and the twin module's bytes are frozen goldens
no refactor should go near.
"""

from __future__ import annotations

import io
import re
import struct
import zipfile
import zlib
from pathlib import Path

# ------------------------------------------------------------- minimal PDF

_PINNED_ZIP_DATE = (2026, 1, 1, 0, 0, 0)


def _pdf_escape(line: str) -> bytes:
    escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    return escaped.encode("latin-1")


def _text_block(lines: list[str], x: int, y: int) -> bytes:
    parts = [b"BT\n/F1 11 Tf\n14 TL\n%d %d Td\n" % (x, y)]
    for i, line in enumerate(lines):
        if i:
            parts.append(b"T*\n")
        parts.append(b"(" + _pdf_escape(line) + b") Tj\n")
    parts.append(b"ET\n")
    return b"".join(parts)


def _assemble_pdf(path: Path, streams: list[bytes]) -> Path:
    """One content stream per page; offsets computed, never guessed."""
    n_pages = len(streams)
    font_num = 3 + 2 * n_pages
    kids = " ".join(f"{3 + 2 * i} 0 R" for i in range(n_pages))
    objs: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>".encode("latin-1"),
        font_num: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    for i, stream in enumerate(streams):
        page_num = 3 + 2 * i
        objs[page_num] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {font_num} 0 R >> >> "
            f"/Contents {page_num + 1} 0 R >>"
        ).encode("latin-1")
        objs[page_num + 1] = (
            b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream)
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += b"%d 0 obj\n" % num + objs[num] + b"\nendobj\n"
    xref_pos = len(out)
    out += b"xref\n0 %d\n" % (len(objs) + 1)
    out += b"0000000000 65535 f \n"
    for num in sorted(objs):
        out += b"%010d 00000 n \n" % offsets[num]
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objs) + 1,
        xref_pos,
    )
    path.write_bytes(bytes(out))
    return path


# ------------------------------------------------------- deterministic docx


def _pinned_docx_bytes(document) -> bytes:
    """python-docx content is deterministic; its zip timestamps are not
    (writestr stamps wall clock — the twins' EC-8 lesson), and a comment
    carries its authoring date. Pin both."""
    buf = io.BytesIO()
    document.save(buf)
    src = zipfile.ZipFile(io.BytesIO(buf.getvalue()))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in sorted(src.namelist()):
            blob = src.read(name)
            if name == "word/comments.xml":
                blob = re.sub(
                    rb'w:date="[^"]*"', b'w:date="2026-01-01T00:00:00Z"', blob
                )
            info = zipfile.ZipInfo(name, date_time=_PINNED_ZIP_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            zout.writestr(info, blob)
    return out.getvalue()


def _shade(cell, fill_hex: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _fill_table(table, grid: list[list[str]]) -> None:
    for r, row in enumerate(grid):
        for c, text in enumerate(row):
            table.rows[r].cells[c].text = text


# ----------------------------------------------------------- corpus content
# Grids below are the READ-BACK view (merged positions repeat their text);
# evals/extraction-gate/ground_truth.json duplicates them as the answer key.

WRITABLE_FILL = "9DC3E6"  # the blue-cells-only convention (spec §A2.3)
LOCKED_FILL = "D9D9D9"
HEADER_FILL = "1F4E79"

COMPLEX_BANNER = "Attachment B - Implementation Pricing (USD)"
COMPLEX_GRID = [
    [COMPLEX_BANNER] * 6,
    ["Workstream", "Module", "Year 1", "Year 2", "Year 3", "Total"],
    ["Core Platform", "Finance", "420,000", "380,000", "120,000", "920,000"],
    ["Core Platform", "Supply Chain", "310,000", "240,000", "90,000", "640,000"],
    ["People", "Human Capital", "280,000", "210,000", "80,000", "570,000"],
    ["Delivery", "Integration Services", "150,000", "140,000", "60,000", "350,000"],
    ["Program Total", "Program Total", "1,160,000", "970,000", "350,000", "2,480,000"],
]
# cell pairs that must be one underlying cell after merging
COMPLEX_MERGES = [([0, 0], [0, 5]), ([2, 0], [3, 0]), ([6, 0], [6, 1])]

SIMPLE_GRID = [
    ["Milestone", "Due", "Owner"],
    ["Discovery complete", "October 2026", "Vendor"],
    ["Finance go-live", "July 2027", "Joint"],
    ["Hypercare exit", "October 2027", "Northwind"],
]

FORM_INSTRUCTION = (
    "Complete the blue answer cells only. Gray cells are locked by the "
    "Northwind procurement office."
)
FORM_GRID = [
    [FORM_INSTRUCTION, FORM_INSTRUCTION],
    ["Vendor legal name", ""],
    ["Years delivering ERP implementations", ""],
    ["Proposed engagement lead", ""],
    ["Total fixed fee (USD)", ""],
]
FORM_MERGES = [([0, 0], [0, 1])]
FORM_COMMENT = "Blue cells only. The fee cell is locked pending Attachment B."

CONTROL_HIERARCHY = [
    [1, "Northwind Regional Health - ERP Implementation Response"],
    [1, "1. Executive Summary"],
    [2, "1.1 Understanding of Objectives"],
    [2, "1.2 Delivery Approach"],
    [3, "1.2.1 Phase Plan"],
    [1, "2. Data Migration Approach"],
    [2, "2.1 Validation Strategy"],
]

MULTICOL_TITLE = "Northwind Regional Health - Conditions of Solicitation"
READING_ORDER = [f"RO-{i:02d} Condition clause number {i} applies." for i in range(1, 13)]

SCANNED_LINES = [
    "NORTHWIND REGIONAL HEALTH",
    "Request for Proposals: ERP Implementation Services",
    "RFP No. NRH-2026-014 (scanned copy)",
    "Proposals are due no later than August 29, 2026.",
    "Submit through the Northwind procurement portal.",
]


# ---------------------------------------------------------------- builders


def build_complex_table_docx(path: Path) -> Path:
    import docx

    doc = docx.Document()
    doc.add_heading("Northwind Regional Health - Solicitation NRH-2026-014", level=1)
    doc.add_heading("Attachment B - Pricing", level=2)
    doc.add_paragraph("Pricing must follow the structure below without alteration.")

    table = doc.add_table(rows=7, cols=6)
    table.style = "Table Grid"
    for a, b in COMPLEX_MERGES:
        table.rows[a[0]].cells[a[1]].merge(table.rows[b[0]].cells[b[1]])
    _fill_table(table, COMPLEX_GRID)
    _shade(table.rows[0].cells[0], HEADER_FILL)
    for cell in table.rows[1].cells:
        _shade(cell, HEADER_FILL)

    doc.add_heading("Attachment C - Milestones", level=2)
    simple = doc.add_table(rows=len(SIMPLE_GRID), cols=3)
    simple.style = "Table Grid"
    _fill_table(simple, SIMPLE_GRID)

    path.write_bytes(_pinned_docx_bytes(doc))
    return path


def build_form_docx(path: Path) -> Path:
    import docx

    doc = docx.Document()
    doc.add_heading("Form B - Vendor Response Form", level=1)
    table = doc.add_table(rows=len(FORM_GRID), cols=2)
    table.style = "Table Grid"
    for a, b in FORM_MERGES:
        table.rows[a[0]].cells[a[1]].merge(table.rows[b[0]].cells[b[1]])
    _fill_table(table, FORM_GRID)
    _shade(table.rows[0].cells[0], HEADER_FILL)
    for r in (1, 2, 3):
        _shade(table.rows[r].cells[1], WRITABLE_FILL)
    _shade(table.rows[4].cells[1], LOCKED_FILL)

    anchor = table.rows[1].cells[1].paragraphs[0].add_run("")
    doc.add_comment(runs=[anchor], text=FORM_COMMENT, author="Procurement", initials="NP")

    path.write_bytes(_pinned_docx_bytes(doc))
    return path


def build_control_docx(path: Path) -> Path:
    import docx

    doc = docx.Document()
    for level, text in CONTROL_HIERARCHY:
        doc.add_heading(text, level=level)
        doc.add_paragraph(
            "Synthetic narrative for the section above: approach, staffing, "
            "and evidence drawn from the Northwind engagement history."
        )
    path.write_bytes(_pinned_docx_bytes(doc))
    return path


def build_table_pdf(path: Path) -> Path:
    """CT-1 as a RULED PDF table: per-logical-cell rectangles + positioned
    text, merges drawn as spanning rects. This is the corpus's real PDF
    table surface — the fabrication two-path diff (deterministic
    TableFormer vs VLM) and the PDF cell-accuracy score both run here;
    the DOCX tables parse XML in both modes and cannot exercise either."""
    x0, y0, col_w, row_h = 72, 700, 78, 24
    continuations = set()
    spans = {}
    for (ar, ac), (br, bc) in COMPLEX_MERGES:
        spans[(ar, ac)] = (br - ar + 1, bc - ac + 1)
        for r in range(ar, br + 1):
            for c in range(ac, bc + 1):
                if (r, c) != (ar, ac):
                    continuations.add((r, c))

    parts = [b"0.5 w\n"]
    for r, row in enumerate(COMPLEX_GRID):
        for c, text in enumerate(row):
            if (r, c) in continuations:
                continue
            rs, cs = spans.get((r, c), (1, 1))
            x, y_top = x0 + c * col_w, y0 - r * row_h
            w, h = cs * col_w, rs * row_h
            parts.append(b"%d %d %d %d re S\n" % (x, y_top - h, w, h))
            parts.append(
                b"BT\n/F1 7 Tf\n%d %d Td\n(%s) Tj\nET\n"
                % (x + 2, y_top - h + (h // 2) - 3, _pdf_escape(text))
            )
    header = _text_block(["Attachment B (ruled table rendering)"], 72, 740)
    return _assemble_pdf(path, [header + b"".join(parts)])


def build_multicolumn_pdf(path: Path) -> Path:
    stream = (
        _text_block([MULTICOL_TITLE], 72, 756)
        + _text_block(READING_ORDER[:6], 72, 700)
        + _text_block(READING_ORDER[6:], 320, 700)
    )
    return _assemble_pdf(path, [stream])


def build_scanned_pdf(path: Path) -> Path:
    """Image-only pages: rendered text, slight rotation, no text layer —
    the scanned-poor-quality stand-in. PIL ships with docling, so this
    builder runs in the gate container only (its trait test rides the
    container-only roster)."""
    from PIL import Image, ImageDraw, ImageFont

    # ~36px glyphs at 1700px page width ≈ 12pt at 150dpi — a realistic
    # office scan. The unsized PIL default (~11px, then halved by the
    # downscale) was OCR-illegible in a way no real scan is (run-3
    # finding); "poor quality" stays in the rotation + downscale.
    font = ImageFont.load_default(size=36)
    pages = []
    for page_no in (1, 2):
        img = Image.new("L", (1700, 2200), 245)
        draw = ImageDraw.Draw(img)
        y = 200
        for line in SCANNED_LINES:
            draw.text(
                (160, y), line if page_no == 1 else f"{line} (page 2)",
                fill=20, font=font,
            )
            y += 90
        pages.append(img.rotate(0.6, expand=False, fillcolor=245).resize((850, 1100)))
    buf = io.BytesIO()
    pages[0].save(buf, "PDF", save_all=True, append_images=pages[1:])
    path.write_bytes(buf.getvalue())
    return path


def build_truncated_pdf(path: Path, source: Path) -> Path:
    data = source.read_bytes()
    path.write_bytes(data[: int(len(data) * 0.55)])
    return path


def build_encrypted_pdf(path: Path, source: Path) -> Path:
    from pypdf import PdfReader, PdfWriter

    writer = PdfWriter()
    writer.append(PdfReader(source))
    writer.encrypt("nrh-locked")
    with path.open("wb") as fh:
        writer.write(fh)
    return path


def build_500_page_pdf(path: Path) -> Path:
    streams = [
        _text_block(
            [
                f"Northwind throughput corpus - page {i} of 500",
                "Body line one for parse timing.",
                "Body line two for parse timing.",
            ],
            72,
            756,
        )
        for i in range(1, 501)
    ]
    return _assemble_pdf(path, streams)


# ------------------------------------------------------------- media (C11)


def _tiny_png(width: int, height: int, bands: list[tuple[int, int, int]]) -> bytes:
    """Stdlib PNG: vertical color bands, RGB8. A committed PNG would fail
    the tripwire's opaque-binary closure; a generated one never lands."""

    def chunk(kind: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + kind
            + data
            + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
        )

    band_w = max(1, width // len(bands))
    raw = b""
    for _y in range(height):
        row = b"\x00"
        for x in range(width):
            row += bytes(bands[min(x // band_w, len(bands) - 1)])
        raw += row
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


LOGO_BANDS = [(31, 78, 121), (157, 195, 230), (191, 191, 191)]
SIGNATURE_BANDS = [(250, 250, 250), (40, 40, 60), (250, 250, 250)]


def build_logo_docx(path: Path) -> Path:
    import docx
    from docx.shared import Inches

    doc = docx.Document()
    doc.add_picture(io.BytesIO(_tiny_png(120, 40, LOGO_BANDS)), width=Inches(1.5))
    doc.add_heading("Northwind Regional Health - Engagement Letter", level=1)
    doc.add_paragraph("The mark above is the buyer's identifying logo.")
    path.write_bytes(_pinned_docx_bytes(doc))
    return path


def build_signature_docx(path: Path) -> Path:
    import docx
    from docx.shared import Inches

    doc = docx.Document()
    doc.add_heading("Certification", level=1)
    doc.add_paragraph("Signed on behalf of the vendor:")
    doc.add_picture(io.BytesIO(_tiny_png(200, 60, SIGNATURE_BANDS)), width=Inches(2))
    doc.add_paragraph("Authorized signature above this line.")
    path.write_bytes(_pinned_docx_bytes(doc))
    return path


def build_logo_case_docx(path: Path) -> Path:
    """C11 anonymization eval case: the logo twin's content shaped for the
    harness — a doc_meta comment paragraph plus a '## '-level section so
    the scripted segmenter produces cards. The embedded PNG is the point:
    identity the text scan cannot see, which the media flag must carry."""
    import docx
    from docx.shared import Inches

    doc = docx.Document()
    doc.add_paragraph(
        "<!-- client: Northwind Regional Health | descriptor: a regional "
        "health system | pursuit: pur_anon_021 | outcome: won | "
        "date: 2026-02-01 -->"
    )
    doc.add_picture(io.BytesIO(_tiny_png(120, 40, LOGO_BANDS)), width=Inches(1.5))
    doc.add_heading("Engagement Summary", level=2)
    doc.add_paragraph(
        "Northwind Regional Health engaged our team to replace its legacy "
        "ERP platform across the health system. The mark above is the "
        "buyer's identifying logo."
    )
    path.write_bytes(_pinned_docx_bytes(doc))
    return path


def build_signature_case_docx(path: Path) -> Path:
    """C11 anonymization eval case: signature-image variant."""
    import docx
    from docx.shared import Inches

    doc = docx.Document()
    doc.add_paragraph(
        "<!-- client: Northwind Regional Health | descriptor: a regional "
        "health system | pursuit: pur_anon_022 | outcome: won | "
        "date: 2026-02-01 -->"
    )
    doc.add_heading("Certification", level=2)
    doc.add_paragraph(
        "Signed on behalf of the vendor for the Northwind Regional Health "
        "engagement:"
    )
    doc.add_picture(io.BytesIO(_tiny_png(200, 60, SIGNATURE_BANDS)), width=Inches(2))
    doc.add_paragraph("Authorized signature above this line.")
    path.write_bytes(_pinned_docx_bytes(doc))
    return path


# ---------------------------------------------------------------- assembly


def build_corpus(workdir: Path) -> dict[str, Path]:
    """Everything buildable without container-only deps. The harness (C5)
    adds scanned-twin.pdf via build_scanned_pdf explicitly — PIL exists
    only where docling does."""
    workdir = Path(workdir)
    workdir.mkdir(parents=True, exist_ok=True)
    paths = {
        "complex-tables-twin.docx": build_complex_table_docx(
            workdir / "complex-tables-twin.docx"
        ),
        "response-form-twin.docx": build_form_docx(workdir / "response-form-twin.docx"),
        "firm-control-twin.docx": build_control_docx(workdir / "firm-control-twin.docx"),
        "multicolumn-twin.pdf": build_multicolumn_pdf(workdir / "multicolumn-twin.pdf"),
        "table-pdf-twin.pdf": build_table_pdf(workdir / "table-pdf-twin.pdf"),
        "500-page-twin.pdf": build_500_page_pdf(workdir / "500-page-twin.pdf"),
        "logo-twin.docx": build_logo_docx(workdir / "logo-twin.docx"),
        "signature-twin.docx": build_signature_docx(workdir / "signature-twin.docx"),
    }
    paths["truncated-twin.pdf"] = build_truncated_pdf(
        workdir / "truncated-twin.pdf", paths["multicolumn-twin.pdf"]
    )
    paths["encrypted-twin.pdf"] = build_encrypted_pdf(
        workdir / "encrypted-twin.pdf", paths["multicolumn-twin.pdf"]
    )
    return paths
