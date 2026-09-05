"""Write-back into the buyer's own DOCX (P16/C8) — the docx twin of
writeback.py, same contract, same posture: it touches a COPY under
exports/writeback/, never the inbox original; preview-then-confirm
(S7) with the facts RE-DERIVED server-side; only prose+authored slots
the frozen plan names, carrying drafted prose, are written; everything
else is a RECORDED refusal.

Docx specifics:
- Answer cells are re-derived from the digest-bound source via
  question_cell_map() (the frozen slot schema carries no table
  addresses — the source file itself is the address book, and the
  digest binding guarantees it is the SAME file the slots came from).
- A prose slot with no addressable cell (a mandated outline section)
  is `refused_no_cell`: its prose ships via the rendered response
  document — recorded, never silently dropped.
- The firm template is NOT this lane: firm_default containers refuse
  (in-place template fill is P17 — the owner's kickoff call, B73§2).
- Round-trip safety is a DOCUMENT-MODEL diff, never bytes (v1's rule:
  the zip re-save is never byte-stable): every paragraph's text+style
  and every table cell outside the intended targets must be identical,
  verified after the write and refused loudly on drift.
"""

import hashlib
import shutil
from pathlib import Path

from docx import Document

from engine.assembly.hygiene import firm_identity, stamp_docx
from engine.contracts import ContractError
from engine.structure.docx_buyer import question_cell_map

WRITEBACK_DIR = "exports/writeback"
FACTS_NAME = "exports/docx-writeback-facts.json"


def _writable(slot: dict) -> tuple[bool, str]:
    shape = slot.get("response_shape")
    if shape != "prose":
        return False, (f"response_shape {shape!r} never takes model prose "
                       "(B28(4) shape-skip: the drafter structurally "
                       "cannot emit it)")
    if slot.get("fill_type") not in (None, "authored"):
        return False, (f"fill_type {slot.get('fill_type')!r} is not the "
                       "authored lane")
    return True, ""


def _docx_source(pursuit, container,
                 binding: dict | None = None) -> tuple[Path, str]:
    """The buyer docx THIS record covers, located by digest — plus the
    merge prefix its slot ids carry ('' when the container is
    single-file). With a binding (P18/C5), exactly that file; without
    one, the container must declare exactly one docx — a multi-docx
    container is served per binding (declared_deliverables), which is
    how B74§3g's recorded limitation retired."""
    if container.get("source_mode") == "firm_default":
        raise ContractError(
            "the slots are the FIRM TEMPLATE's (firm_default) — this lane "
            "fills buyer questionnaires; the firm template fills through "
            "engine.assembly.template_fill (P17/C9, live), which the web "
            "writeback dispatcher routes to by source_mode")
    wanted: list[tuple[str, str]] = []  # (sha256, prefix)
    if binding:
        wanted.append((binding["source_sha256"], binding["prefix"]))
    else:
        sources = container.get("sources")
        if sources:
            docx_entries = [(i, e) for i, e in enumerate(sources)
                            if e["file"].lower().endswith(".docx")]
            if len(docx_entries) > 1:
                raise ContractError(
                    "this container declares several docx sources — "
                    "each fills through its OWN binding "
                    "(declared_deliverables, P18/C5); a bare call "
                    "cannot say which file it means")
            for i, entry in docx_entries:
                wanted.append((entry["source_sha256"], f"f{i:02d}-"))
        else:
            wanted.append((container["source_sha256"], ""))
    matches: list[tuple[Path, str]] = []
    for digest, prefix in wanted:
        for candidate in sorted((pursuit.root / "inbox").glob("*.docx")):
            if hashlib.sha256(
                    candidate.read_bytes()).hexdigest() == digest:
                matches.append((candidate, prefix))
    if not matches:
        raise ContractError(
            "no inbox docx matches the slots container's digest — docx "
            "write-back only fills the EXACT file the slots were parsed "
            "from")
    return matches[0]


def compute_docx_facts(pursuit, *, at: str, confirmed_by: str,
                       binding: dict | None = None) -> dict:
    """binding (P18/C5): one entry from declared_deliverables — this
    record covers exactly that file's slots (the cross-file filter below
    is the per-binding contract, not a silent drop: the other files'
    slots belong to the other files' records, and the union covers every
    slot — the named coverage test)."""
    frozen_plan = pursuit.read_frozen("pursuit_plan")
    container = pursuit.read_artifact(
        frozen_plan.get("slots_ref", "slots.json"))
    envelope = pursuit.read_artifact("drafts/draft.json")
    source, prefix = _docx_source(pursuit, container, binding)
    cell_map = question_cell_map(source)

    prose_by_slot = {a["slot_id"]: a
                     for e in envelope.get("sections", [])
                     for a in e.get("answers", [])}
    planned_slots = {sid for s in frozen_plan.get("sections", [])
                     for sid in s.get("slot_ids", [])}

    cells = []
    for slot in container["slots"]:
        locator = slot.get("source_locator", {})
        if locator.get("file") != source.name or slot.get("is_header"):
            continue
        slot_id = slot["slot_id"]
        row = {"slot_id": slot_id}
        if slot.get("ref_id"):
            row["ref_id"] = slot["ref_id"]
        if locator.get("docx_anchor"):
            row["docx_anchor"] = locator["docx_anchor"]
        address = cell_map.get(
            slot_id[len(prefix):] if prefix and slot_id.startswith(prefix)
            else slot_id)
        writable, why = _writable(slot)
        answer = prose_by_slot.get(slot_id, {})
        if slot_id not in planned_slots:
            row["decision"] = "refused_unnamed"
            row["reason"] = ("the frozen plan names no section for this "
                            "slot — T6: write-back may only touch what "
                            "the pursuit plan names")
        elif not writable:
            row["decision"] = "refused_shape"
            row["reason"] = why
        elif address is None:
            row["decision"] = "refused_no_cell"
            row["reason"] = ("no addressable answer cell in the source — "
                            "this section's prose ships via the rendered "
                            "response document (P16/C8)")
        elif answer.get("status") == "drafted" and answer.get("prose"):
            row.update(address)
            row["decision"] = "written"
            row["reason"] = "drafted prose from the validated envelope"
            row["after"] = answer["prose"]
        else:
            row.update(address)
            row["decision"] = "empty_no_prose"
            row["reason"] = (f"slot status "
                            f"{answer.get('status', 'undrafted')!r} — the "
                            "cell stays honestly empty for the human")
        cells.append(row)

    document = Document(str(source))
    for row in cells:
        if row["decision"] == "written":
            cell = (document.tables[row["table_index"]]
                    .rows[row["row"]].cells[row["column"]])
            row["before"] = cell.text

    return {
        "pursuit_id": pursuit.pursuit_id,
        "plan_sha256": envelope["plan_sha256"],
        "draft_sha256": hashlib.sha256(
            (pursuit.root / "drafts" / "draft.json").read_bytes()
        ).hexdigest(),
        "revision_n": envelope["revision_n"],
        "confirmed_by": confirmed_by,
        "at": at,
        "source_file": str(source.relative_to(pursuit.root)),
        "output_file": f"{WRITEBACK_DIR}/{source.name}",
        "cells": cells,
    }


def preview_docx_writeback(pursuit, *, at: str,
                           binding: dict | None = None) -> dict:
    """The S7 preview: the facts WITHOUT the write."""
    return compute_docx_facts(pursuit, at=at, confirmed_by="(unconfirmed)",
                              binding=binding)


def _document_model(path: Path) -> tuple[list, list]:
    doc = Document(str(path))
    paragraphs = [(p.text, p.style.name if p.style else "")
                  for p in doc.paragraphs]
    tables = [[[cell.text for cell in row.cells] for row in table.rows]
              for table in doc.tables]
    return paragraphs, tables


def _assert_roundtrip(source: Path, output: Path,
                      intended: set[tuple[int, int, int]]) -> None:
    """v1's rule, docx edition: prove 'we changed nothing else' on the
    document MODEL. A drifted paragraph or an unintended cell is a
    loud refusal, never a shipped surprise."""
    src_paras, src_tables = _document_model(source)
    out_paras, out_tables = _document_model(output)
    if src_paras != out_paras:
        raise ContractError(
            "docx write-back drifted paragraph content or styles outside "
            "the intended cells — refusing to hand back a changed "
            "document")
    if len(src_tables) != len(out_tables):
        raise ContractError("docx write-back changed the table count")
    for t, (src_t, out_t) in enumerate(zip(src_tables, out_tables)):
        for r, (src_row, out_row) in enumerate(zip(src_t, out_t)):
            for c, (src_cell, out_cell) in enumerate(zip(src_row, out_row)):
                if src_cell != out_cell and (t, r, c) not in intended:
                    raise ContractError(
                        f"docx write-back drifted table {t} r{r}c{c} — "
                        "outside every intended target")


def run_docx_writeback(pursuit, log, *, at: str, confirmed_by: str,
                       binding: dict | None = None) -> dict:
    """The confirmed write: re-derives the facts server-side, copies the
    buyer file, assigns ONLY the written cells, proves the round-trip,
    records everything."""
    facts = compute_docx_facts(pursuit, at=at, confirmed_by=confirmed_by,
                               binding=binding)
    source = pursuit.root / facts["source_file"]
    output = pursuit.root / facts["output_file"]
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, output)
    document = Document(str(output))
    intended = set()
    for row in facts["cells"]:
        if row["decision"] == "written":
            address = (row["table_index"], row["row"], row["column"])
            intended.add(address)
            (document.tables[address[0]].rows[address[1]]
             .cells[address[2]]).text = row["after"]
    # P3-15: the buyer's form keeps the buyer's author; the firm is the
    # last modifier, and no generator string rides along.
    stamp_docx(document, firm=firm_identity(pursuit.root.parent), at=at,
               title=document.core_properties.title or "",
               owned_by_firm=False)
    document.save(output)
    _assert_roundtrip(source, output, intended)
    facts_path = pursuit.write_artifact(
        "writeback_facts", facts,
        name=binding["facts_name"] if binding else FACTS_NAME)
    log.emit("artifact", stage="write_back", artifact={
        "kind": "writeback_facts", "path": str(facts_path),
        "revision_n": facts["revision_n"],
        "sha256": hashlib.sha256(facts_path.read_bytes()).hexdigest()})
    log.emit("artifact", stage="write_back", artifact={
        "kind": "write_back_file", "path": str(output),
        "revision_n": facts["revision_n"],
        "sha256": hashlib.sha256(output.read_bytes()).hexdigest()})
    return facts
