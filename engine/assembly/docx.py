"""Export-to-Word (B37/D20): two lanes split ON DISK, because "which of
these does the buyer get?" must be answerable by looking (v1's two-
literal-headings lesson) — `exports/submission/` is To the buyer,
`exports/review/` is Internal — do not send.

The SUBMISSION export refuses while packaging is blocked or a drafted-
owed section still pends: an export door that opens under a block would
be the control leaking at the exit. The REVIEW export always renders —
flags, pends, and findings visible — because the internal reader needs
the whole truth. python-docx does the rendering (no hand-rolled DOCX,
spec CLAUDE rule); neutral styling only (D31): the buyer-facing document
carries the pursuit's buyer name and section titles, never engine
vocabulary; real branding is the ≤A6 asset swap."""

import hashlib

from docx import Document

from engine.assembly.bindings import owed_pends as _owed_pends
from engine.contracts import ContractError

SUBMISSION_NAME = "exports/submission/response.docx"
REVIEW_NAME = "exports/review/annotated-review.docx"


def _load(pursuit):
    envelope = pursuit.read_artifact("drafts/draft.json")
    annotated = pursuit.read_artifact("drafts/annotated-draft.json")
    brief = pursuit.read_frozen("bid_brief")  # verified (P0-2)
    return envelope, annotated, brief


def _title(brief) -> str:
    buyer = brief.get("buyer", {}).get("name", "")
    return (f"Response — {buyer}" if buyer else "Response")


def _refuse_firm_default(pursuit) -> None:
    """P17/C10 (B75§1d, the owner's call): for a firm-template pursuit the
    FILLED template IS the to-the-buyer document — one submission
    document per pursuit, so the generated render refuses with the
    pointer instead of minting a competitor."""
    try:
        frozen = pursuit.read_frozen("pursuit_plan")
        container = pursuit.read_artifact(
            frozen.get("slots_ref", "slots.json"))
    except FileNotFoundError:
        return
    if container.get("source_mode") == "firm_default":
        raise ContractError(
            "this pursuit plans against the firm template — its "
            "submission document is the FILLED template (preview/confirm "
            "the fill via the writeback lanes), not a generated render "
            "(B75§1d)")


def render_submission(pursuit, log, *, at: str) -> str:
    _refuse_firm_default(pursuit)
    envelope, annotated, brief = _load(pursuit)
    packaging = annotated.get("packaging", {})
    if packaging.get("blocked"):
        raise ContractError(
            f"packaging is BLOCKED ({packaging.get('tier1_blocks', 0)} "
            "tier-1 block(s)) — the submission export never opens under "
            "a block; waive or revise first")
    pends = _owed_pends(envelope)
    if pends:
        raise ContractError(
            "drafted-owed section(s) still pend gap dispositions: "
            + ", ".join(pends) + " — a submission with silent holes "
            "would misrepresent the response")
    doc = Document()
    doc.add_heading(_title(brief), level=0)
    appendices = []
    for entry in envelope.get("sections", []):
        if entry.get("status") != "drafted":
            continue  # approved omissions simply do not appear
        doc.add_heading(entry.get("title", entry["section_id"]), level=1)
        for answer in entry.get("answers", []):
            if not answer.get("prose"):
                continue
            if answer.get("omission_stated"):
                doc.add_paragraph(answer["prose"])
                continue
            appendix = answer.get("appendix_routed")
            if appendix:
                appendices.append((entry, answer))
                doc.add_paragraph(f"See Appendix — {answer.get('ref_id') or answer['slot_id']}.")
            else:
                doc.add_paragraph(answer["prose"])
        if entry.get("prose"):
            doc.add_paragraph(entry["prose"])
    if appendices:
        doc.add_heading("Appendices", level=1)
        for entry, answer in appendices:
            doc.add_heading(
                f"Appendix — {answer.get('ref_id') or answer['slot_id']}",
                level=2)
            doc.add_paragraph(answer["prose"])
    path = pursuit.root / SUBMISSION_NAME
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    log.emit("artifact", stage="write_back", artifact={
        "kind": "export", "path": str(path),
        "revision_n": envelope["revision_n"],
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest()})
    return SUBMISSION_NAME


def render_review(pursuit, log, *, at: str) -> str:
    envelope, annotated, brief = _load(pursuit)
    marks = {s["section_id"]: s for s in annotated.get("sections", [])}
    doc = Document()
    doc.add_heading(_title(brief), level=0)
    doc.add_paragraph("Internal — do not send. Flags, pends, and "
                      "validation findings are visible in this copy.")
    packaging = annotated.get("packaging", {})
    doc.add_paragraph(
        f"Packaging: {'BLOCKED' if packaging.get('blocked') else 'clear'}"
        f" — {packaging.get('tier1_blocks', 0)} tier-1 block(s), "
        f"{packaging.get('waived', 0)} waived. Revision "
        f"{envelope.get('revision_n')}.")
    for entry in envelope.get("sections", []):
        doc.add_heading(entry.get("title", entry["section_id"]), level=1)
        if entry.get("status") != "drafted":
            doc.add_paragraph(
                f"[{entry.get('status', 'undrafted').upper()}: "
                f"{entry.get('reason', 'no prose')}]")
        for answer in entry.get("answers", []):
            if answer.get("prose"):
                doc.add_paragraph(answer["prose"])
            elif answer.get("status") == "awaiting_disposition":
                doc.add_paragraph(
                    f"[PENDING {answer['slot_id']}: "
                    f"{answer.get('reason', 'awaiting disposition')}]")
        if entry.get("prose"):
            doc.add_paragraph(entry["prose"])
        section_marks = marks.get(entry["section_id"], {})
        for claim in section_marks.get("claims", []):
            if claim.get("disposition") in ("block", "flag", "waived"):
                doc.add_paragraph(
                    f"• claim {claim.get('status')}: "
                    f"{claim.get('text', '')[:120]}", style="List Bullet")
        for finding in section_marks.get("findings", []):
            doc.add_paragraph(
                f"• {finding.get('check')}/{finding.get('rule')}: "
                f"{finding.get('message', '')[:160]}", style="List Bullet")
    path = pursuit.root / REVIEW_NAME
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    log.emit("artifact", stage="write_back", artifact={
        "kind": "export", "path": str(path),
        "revision_n": envelope["revision_n"],
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest()})
    return REVIEW_NAME
