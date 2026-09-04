"""In-place firm-template fill (P17/C9 — B73§2's deferral closed; P26a
item 1 — P1-27: the authoring scaffolding leaves, the hand-completion
record lands, and the buyer copy is withheld until nothing remains).

The firm_default lane's exit door: for a Path-B pursuit planned against
the firm's own template, the drafted section prose FILLS the template in
place — REPLACEMENT, not insertion (the v1 oracle's law, reimplemented):
each answered section's prose takes its "▸ WHAT TO INCLUDE" guidance
box's place, and the "[ Replace with … ]" placeholder table goes too
(the v1 gap FIXED — v1 shipped documents still carrying placeholders).
The shapes the engine never drafts — the proposal-metadata record, the
pricing grid, the case block, the inline bracketed line — render from
the hand-completion record (exports/hand-fill.json, entered through the
PUT hand-fill door) and are otherwise reported as owed.

Two copies from one fill (P1-27). The WORKING copy (exports/review/) is
written every time: the template's own authoring scaffolding — its
title, the "How to Use This Template" heading, the instruction
paragraph, everything before the first numbered section except the
metadata table — is stripped; hand values are rendered; a section a
human still drafts keeps its guidance IDENTICAL (the hand-drafter needs
it). The BUYER copy (exports/submission/response.docx — B75§1d: it
replaces the generated render) is written ONLY when nothing remains:
no undrafted section, no owed hand slot, no unnamed anchor. Until then
the bundle records the buyer deliverable as refused with the itemized
reason, so the downloads door never hands a buyer a document that says
"Replace with the drafted section".

Source binding: the template at config/templates/, verified against the
reference_sha256 the path_b_outline checkpoint recorded — the fill only
touches the EXACT template the plan was built against (B74§3b's
digest-bound rule); placement is re-derived by REPARSE + body walk,
never persisted. Prose lands exactly as render_submission would land
it — one exit-door posture, no divergent hygiene.

The verifier is fill-shaped (B75§3c): a body-order stream diff — the
output must equal the source stream with exactly the intended elements
removed, replaced-by-paragraphs, cell-written, or text-replaced, and
NOTHING else different. v1 never round-tripped its fill at all; this
lane refuses to hand back a document it cannot prove.
"""

import hashlib
import re
import shutil
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from engine.assembly.hand_fill import (
    case_block_text,
    completeness,
    hand_slots,
    read_hand_fill,
)
from engine.contracts import ContractError, validate
from engine.planning.plan import REFERENCE_DEFAULT
from engine.structure.docx_default import parse_default_template

FACTS_NAME = "exports/template-fill-facts.json"
OUTPUT_NAME = "exports/submission/response.docx"
WORKING_NAME = "exports/review/response-working.docx"

_P_TAG = "}p"
_TBL_TAG = "}tbl"
_NUMBERED_H1 = re.compile(r"^(\d+)\.\s+")
_BRACKET_SPAN = re.compile(r"\[[^\]]*\]")


def _template_source(pursuit):
    frozen = pursuit.read_frozen("pursuit_plan")
    container = pursuit.read_artifact(frozen.get("slots_ref", "slots.json"))
    if container.get("source_mode") != "firm_default":
        raise ContractError(
            "template fill is the firm_default lane — a buyer-provided "
            "target ships through write-back, not the firm template")
    ckpt = pursuit.checkpoint_payload("path_b_outline")
    ref_sha = ckpt.get("reference_sha256")
    template = REFERENCE_DEFAULT
    actual = hashlib.sha256(template.read_bytes()).hexdigest()
    if actual != ref_sha:
        raise ContractError(
            "the firm template drifted since planning — refusing to fill "
            "a template the frozen plan was not built against (re-run "
            "planning against the current template)")
    return template, ref_sha, frozen, container


def _hand_values(pursuit, ref_sha: str) -> dict:
    """The hand-completion record's values, bound to THIS template."""
    record = read_hand_fill(pursuit)
    if record is None:
        return {}
    if record.get("template_sha256") != ref_sha:
        raise ContractError(
            "the hand-completion record was entered against a different "
            "firm template — re-enter the values through the hand-fill "
            "door (PUT …/writeback/hand-fill) before filling")
    return record.get("values", {})


def _body_stream(doc) -> list:
    """The document body in order: ('p', text, style) | ('t', cell rows).
    The comparison currency of the fill verifier; an element's position
    in this stream is its identity for the intended-change set."""
    stream = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith(_P_TAG):
            para = Paragraph(child, doc)
            stream.append(("p", para.text,
                           para.style.name if para.style else ""))
        elif child.tag.endswith(_TBL_TAG):
            table = Table(child, doc)
            stream.append(("t", [[cell.text for cell in row.cells]
                                 for row in table.rows]))
    return stream


def _delivered_paragraphs(prose: str) -> list[str]:
    return [part.strip() for part in prose.split("\n\n") if part.strip()]


def _plan_bindings(frozen) -> dict[str, str]:
    """slot_id -> plan section_id, through the slot_ids the architect's
    sections inherited from the template (based_on, P16/C7)."""
    bindings: dict[str, str] = {}
    for section in frozen.get("sections", []):
        for slot_id in section.get("slot_ids", []):
            bindings[slot_id] = section["section_id"]
    return bindings


def _fields_written(slot: dict, value) -> list[str]:
    shape = slot.get("response_shape")
    if value is None:
        return []
    if shape == "record":
        return sorted(k for k, v in value.items() if v)
    if shape == "table":
        return sorted({k for row in value for k, v in row.items() if v})
    return ["value"]


def _front_matter_texts(template: Path) -> list[str]:
    """What the fill strips: every body element before the first numbered
    Heading 1 except the metadata table — by text, for the record."""
    doc = Document(str(template))
    removed = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith(_P_TAG):
            para = Paragraph(child, doc)
            style = para.style.name if para.style else ""
            if style == "Heading 1" and _NUMBERED_H1.match(para.text.strip()):
                break
            if para.text.strip():
                removed.append(para.text.strip())
        elif child.tag.endswith(_TBL_TAG):
            table = Table(child, doc)
            first = table.rows[0].cells[0].text.strip() if table.rows else ""
            if first.lower() != "field":
                removed.append(first)
    return removed


def compute_fill_facts(pursuit, *, confirmed_by: str, at: str) -> dict:
    """One row per non-header template slot — filled, filled_by_hand,
    kept_guidance, fill_by_hand, or refused_unnamed; re-derived
    server-side every time (S7: the preview and the run compute the
    same facts)."""
    template, ref_sha, frozen, _container = _template_source(pursuit)
    envelope = pursuit.read_artifact("drafts/draft.json")
    prose_by_section = {
        s["section_id"]: s.get("prose", "")
        for s in envelope.get("sections", [])
        if s.get("status") == "drafted" and s.get("prose")}
    bindings = _plan_bindings(frozen)
    parsed = parse_default_template(template)
    hand_values = _hand_values(pursuit, ref_sha)
    hand_ids = {s["slot_id"] for s in hand_slots(parsed)}

    rows = []
    remaining: list[str] = []
    remaining_by_hand: list[str] = []
    for slot in parsed.slots:
        if slot.get("is_header"):
            continue
        anchor = slot["source_locator"]["docx_anchor"]
        question = slot.get("question_text", "")
        section_id = bindings.get(slot["slot_id"]) or \
            bindings.get(slot.get("parent", ""), "")
        row = {"section_id": section_id or "", "slot_id": slot["slot_id"],
               "docx_anchor": anchor}
        shape = slot.get("response_shape")
        if slot["slot_id"] in hand_ids:
            value = hand_values.get(slot["slot_id"])
            complete, missing = completeness(slot, value)
            row["fields_written"] = _fields_written(slot, value)
            if complete:
                row["decision"] = "filled_by_hand"
                row["reason"] = ("the hand-completion record supplied "
                                 "every field — rendered in place")
            else:
                row["decision"] = "fill_by_hand"
                row["reason"] = (f"a human still owes this "
                                 f"{shape} slot — missing: "
                                 + ", ".join(missing))
                remaining_by_hand.append(
                    f"{slot.get('path', anchor)}: missing "
                    + ", ".join(missing))
        elif shape == "prose" and question.startswith("▸"):
            if not section_id:
                row["decision"] = "refused_unnamed"
                row["reason"] = ("the architect's adapted plan carries no "
                                 "section for this template anchor — "
                                 "guidance kept")
                remaining.append(question)
            elif section_id in prose_by_section:
                row["decision"] = "filled"
                row["reason"] = "drafted section prose replaces the guidance"
                row["paragraphs"] = len(_delivered_paragraphs(
                    prose_by_section[section_id]))
                row["placeholder_removed"] = True
            else:
                row["decision"] = "kept_guidance"
                row["reason"] = ("no drafted prose for this section — its "
                                 "guidance survives identical in the "
                                 "working copy (the honest outcome)")
                remaining.append(question)
        else:
            row["decision"] = "fill_by_hand"
            row["reason"] = (f"no fill exists for this {shape} slot — a "
                             "human completes it")
            remaining_by_hand.append(f"{slot.get('path', anchor)}: no fill "
                                     f"for this shape")
        rows.append(row)

    unnamed = any(r["decision"] == "refused_unnamed" for r in rows)
    facts = {
        "pursuit_id": pursuit.pursuit_id,
        "plan_sha256": pursuit.file_sha256("plan.frozen.json"),
        "draft_sha256": pursuit.file_sha256("drafts/draft.json"),
        "revision_n": int(envelope.get("revision_n", 0)),
        "confirmed_by": confirmed_by,
        "at": at,
        "template_file": str(template),
        "template_sha256": ref_sha,
        "output_file": OUTPUT_NAME,
        "working_copy": WORKING_NAME,
        "buyer_copy_produced": (not remaining and not remaining_by_hand
                                and not unnamed),
        "scaffolding_removed": _front_matter_texts(template),
        "sections": rows,
        "remaining_guidance": remaining,
        "remaining_by_hand": remaining_by_hand,
    }
    validate("template_fill_facts", facts)
    return facts


def preview_template_fill(pursuit, *, at: str) -> dict:
    return compute_fill_facts(pursuit, confirmed_by="(unconfirmed)", at=at)


def withheld_reason(facts: dict) -> str:
    """Why the buyer copy was not written — the bundle's refusal text."""
    items = list(facts.get("remaining_by_hand", []))
    n_guidance = len(facts.get("remaining_guidance", []))
    if n_guidance:
        items.append(f"{n_guidance} section(s) still carry guidance "
                     "(no drafted prose)")
    n_unnamed = sum(1 for r in facts.get("sections", [])
                    if r.get("decision") == "refused_unnamed")
    if n_unnamed:
        items.append(f"{n_unnamed} template anchor(s) the plan names no "
                     "section for")
    return (f"buyer copy withheld — {len(items)} item(s) remain: "
            + "; ".join(items)
            + f"; working copy at {facts.get('working_copy', WORKING_NAME)}")


# --------------------------------------------------------------- rendering


def _record_rows(table_rows: list[list[str]], slot: dict, value: dict,
                 *, drop_header: bool) -> list[list[str]]:
    rows = [list(r) for r in table_rows]
    for field in slot.get("response_fields", []):
        loc = field.get("source_locator", {})
        r, c = loc.get("row"), loc.get("column")
        if r is None or c is None:
            continue
        text = (value or {}).get(field["key"])
        if text:
            rows[r][c] = text
    return rows[1:] if drop_header else rows


def _grid_rows(table_rows: list[list[str]], slot: dict,
               entries: list[dict]) -> list[list[str]]:
    header = list(table_rows[0])
    width = len(header)
    column_of = {f["key"]: f.get("source_locator", {}).get("column")
                 for f in slot.get("response_fields", [])}
    out = [header]
    for entry in entries or []:
        cells = [""] * width
        for key, text in entry.items():
            c = column_of.get(key)
            if c is not None and c < width and text:
                cells[c] = text
        out.append(cells)
    return out


_case_block_text = case_block_text  # one renderer (moved to hand_fill, P26c)


def _set_table(table, rows: list[list[str]]) -> None:
    """Make the python-docx table read back exactly `rows`: rewrite
    cells in place, add rows past the template's, drop the surplus."""
    while len(table.rows) < len(rows):
        table.add_row()
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            if table.rows[r].cells[c].text != text:
                table.rows[r].cells[c].text = text
    for surplus in list(table.rows)[len(rows):]:
        table._tbl.remove(surplus._tr)


def _fill_document(doc, facts, prose_by_section, hand_values, parsed,
                   *, buyer: bool) -> dict[int, tuple]:
    """The XML surgery (v1's _replace_box, reimplemented not ported), in
    body order, keyed by stream position: the front matter goes (all but
    the metadata table); a filled section's delivered paragraphs take
    its guidance table's place and its placeholder table goes; a hand
    slot renders its record's values (cells, cell text, or the bracketed
    span); a section a human still drafts keeps guidance + placeholder.
    Returns the intended-change set the verifier replays."""
    filled = {row["docx_anchor"]: row for row in facts["sections"]
              if row["decision"] == "filled"}
    hand_done = {row["slot_id"] for row in facts["sections"]
                 if row["decision"] == "filled_by_hand"}
    question_by_anchor = {
        s["source_locator"]["docx_anchor"]: s.get("question_text", "")
        for s in parsed.slots
        if s.get("response_shape") == "prose"
        and s.get("question_text", "").startswith("▸")}
    hand_by_table = {}   # table_index -> slot, by IDENTITY (value or not)
    inline_by_text = {}
    for slot in hand_slots(parsed):
        if slot.get("response_shape") in ("record", "table"):
            ti = (slot.get("response_fields") or [{}])[0].get(
                "source_locator", {}).get("table_index")
            if ti is not None:
                hand_by_table[ti] = slot
        elif hand_values.get(slot["slot_id"]) is not None:
            inline_by_text[slot.get("question_text", "")] = (
                slot, hand_values[slot["slot_id"]])
    # the guidance box above a hand-completed table goes with it — the
    # parser hands the ▸ text to the TABLE slot (no prose slot exists
    # for those sections), so match on the slot's own question_text
    hand_guidance_texts = {
        slot.get("question_text", "")
        for slot in hand_slots(parsed)
        if slot["slot_id"] in hand_done
        and slot.get("response_shape") == "table"
        and slot.get("question_text", "").startswith("▸")}

    intended: dict[int, tuple] = {}
    k = -1
    tbl_index = -1
    in_front_matter = True
    pending_placeholder = False
    for child in list(doc.element.body.iterchildren()):
        is_p = child.tag.endswith(_P_TAG)
        is_t = child.tag.endswith(_TBL_TAG)
        if not (is_p or is_t):
            continue
        k += 1
        if is_p:
            para = Paragraph(child, doc)
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if in_front_matter:
                if style == "Heading 1" and _NUMBERED_H1.match(text):
                    in_front_matter = False
                else:
                    intended[k] = ("remove",)
                    child.getparent().remove(child)
                    continue
            hit = inline_by_text.get(text)
            if hit is not None:
                _slot, value = hit
                new_text = _BRACKET_SPAN.sub(lambda _m: value,
                                             para.text, count=1)
                para.text = new_text
                intended[k] = ("p_text", new_text)
            continue

        tbl_index += 1
        table = Table(child, doc)
        if not table.rows or not table.rows[0].cells:
            continue
        first_cell = table.rows[0].cells[0].text.strip()
        source_rows = [[cell.text for cell in row.cells]
                       for row in table.rows]
        if tbl_index in hand_by_table:
            slot = hand_by_table[tbl_index]
            value = hand_values.get(slot["slot_id"])
            if value is None:
                pending_placeholder = False
                continue  # a hand slot still owed keeps its template shape
            if slot.get("response_shape") == "record":
                rows = _record_rows(source_rows, slot, value,
                                    drop_header=buyer)
            elif len(slot.get("response_fields", [])) and \
                    len(table.columns) >= 2:
                rows = _grid_rows(source_rows, slot, value)
            else:
                rows = [[_case_block_text(slot, value)]]
            _set_table(table, rows)
            intended[k] = ("cells", rows)
            pending_placeholder = False
            continue
        if in_front_matter:
            intended[k] = ("remove",)
            child.getparent().remove(child)
            continue
        if pending_placeholder:
            pending_placeholder = False
            if first_cell.startswith("["):
                # the placeholder that trailed a filled guidance box
                intended[k] = ("remove",)
                child.getparent().remove(child)
                continue
        match = next((a for a, q in question_by_anchor.items()
                      if q == first_cell and a in filled), None)
        if match is None:
            if first_cell in hand_guidance_texts:
                intended[k] = ("remove",)
                child.getparent().remove(child)
            continue
        row = filled.pop(match)
        parts = _delivered_paragraphs(prose_by_section[row["section_id"]])
        parent = child.getparent()
        position = list(parent).index(child)
        anchor_para = doc.add_paragraph()
        for offset, part in enumerate(parts):
            para = doc.add_paragraph(part)
            parent.insert(position + offset, para._element)
        anchor_para._element.getparent().remove(anchor_para._element)
        parent.remove(child)
        intended[k] = ("replace", parts)
        pending_placeholder = True
    if filled:
        raise ContractError(
            "template fill could not locate the guidance box for: "
            + ", ".join(sorted(filled)) + " — the template and the plan "
            "disagree; refusing a partial fill")
    return intended


def _assert_fill_roundtrip(template: Path, output: Path,
                           intended: dict[int, tuple]) -> None:
    """The fill-shaped verifier (B75§3c): the output body stream must
    equal the source stream with EXACTLY the intended elements removed,
    replaced by their delivered paragraphs, cell-written, or
    text-replaced — untouched paragraphs identical in text and style,
    untouched tables cell-identical."""
    source = Document(str(template))
    expected: list = []
    for k, item in enumerate(_body_stream(source)):
        action = intended.get(k)
        if action is None:
            expected.append(item)
        elif action[0] == "remove":
            continue
        elif action[0] == "replace":
            expected.extend(("p", part, "Normal") for part in action[1])
        elif action[0] == "cells":
            expected.append(("t", action[1]))
        elif action[0] == "p_text":
            expected.append(("p", action[1], item[2]))
        else:
            raise ContractError(f"unknown fill action {action[0]!r}")
    actual = _body_stream(Document(str(output)))
    if actual != expected:
        raise ContractError(
            "template fill drifted outside the intended sections — "
            "refusing to hand back a document that differs from the "
            "declared change set")


def _render(template: Path, target: Path, facts, prose_by_section,
            hand_values, parsed, *, buyer: bool) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, target)
    doc = Document(str(target))
    intended = _fill_document(doc, facts, prose_by_section, hand_values,
                              parsed, buyer=buyer)
    doc.save(str(target))
    _assert_fill_roundtrip(template, target, intended)


def run_template_fill(pursuit, log, *, confirmed_by: str, at: str) -> dict:
    """Confirm-and-run (S7): facts re-derived server-side, the working
    copy always, the buyer copy only when nothing remains, each proven
    by the stream-diff verifier, the facts artifact written and logged."""
    facts = compute_fill_facts(pursuit, confirmed_by=confirmed_by, at=at)
    template = Path(facts["template_file"])
    envelope = pursuit.read_artifact("drafts/draft.json")
    prose_by_section = {
        s["section_id"]: s.get("prose", "")
        for s in envelope.get("sections", [])
        if s.get("status") == "drafted" and s.get("prose")}
    parsed = parse_default_template(template)
    hand_values = _hand_values(pursuit, facts["template_sha256"])

    working = pursuit.root / WORKING_NAME
    _render(template, working, facts, prose_by_section, hand_values,
            parsed, buyer=False)
    output = pursuit.root / OUTPUT_NAME
    if facts["buyer_copy_produced"]:
        _render(template, output, facts, prose_by_section, hand_values,
                parsed, buyer=True)
    elif output.exists():
        output.unlink()  # a stale buyer copy never outlives its facts

    facts_path = pursuit.write_artifact("template_fill_facts", facts,
                                        name=FACTS_NAME)
    log.emit("artifact", stage="write_back", artifact={
        "kind": "template_fill_facts", "path": str(facts_path),
        "revision_n": facts["revision_n"],
        "sha256": hashlib.sha256(facts_path.read_bytes()).hexdigest()})
    log.emit("artifact", stage="write_back", artifact={
        "kind": "export", "path": str(working),
        "revision_n": facts["revision_n"],
        "sha256": hashlib.sha256(working.read_bytes()).hexdigest()})
    if facts["buyer_copy_produced"]:
        log.emit("artifact", stage="write_back", artifact={
            "kind": "export", "path": str(output),
            "revision_n": facts["revision_n"],
            "sha256": hashlib.sha256(output.read_bytes()).hexdigest()})
    return facts
