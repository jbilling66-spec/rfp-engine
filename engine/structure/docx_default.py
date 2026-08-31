"""The firm-default template parser: the firm's own Word template
becomes TargetSlots (P16/C2; source_mode="firm_default" — the enum value
the frozen schema declared with no writer until now).

Reimplemented from the v1 oracle (v1 engine/structure/docx_default.py,
read-only scratchpad clone; reimplement, don't port — B56). The
template's conventions, verified against template-twin.docx:

- Body children INTERLEAVE paragraphs and tables — the walk uses
  body.iterchildren(); python-docx's .paragraphs/.tables lose order.
- Numbering is literal heading text ("1.  Cover Letter", two spaces);
  slot extraction starts at the first Heading 1 matching ^\\d+\\.\\s —
  everything before it is authoring front matter, EXCEPT the proposal-
  metadata table, which becomes ONE record slot (assembly must be able
  to address those cells).
- Tables carry NO named styles; classification is content-based:
  first cell "▸ …" → guidance (held as the next placeholder's
  question_text; stripped from rendered output — the P17 fill lane);
  first cell "[ …" → a prose placeholder slot, unless it carries >=2
  inner "Label: [" fields (the case-study block) which makes it a
  TABLE slot with those fields; multi-column tables with a header row
  → TABLE slots with typed fields.
- "Keep it to one page" in guidance → constraints {brevity: terse,
  flags: [one_page]} (the v1 rule, unchanged).
- docx_anchor is the VERBATIM heading text ("H1" or "H1 > H2") for
  round-trip fidelity; the human path gets cleaned titles.

Slots are plain dicts shaped per the frozen target-slot schema, absent
fields OMITTED never nulled (the v2 writers-omit rule); the caller
validates. Deterministic code end to end; zero model calls.
"""

import hashlib
import re
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from engine.contracts.slots import field_key
from engine.structure.classify import _field_type, _slot
from engine.structure.parse import ParsedWorkbook, StructureError

_NUMBERED_H1 = re.compile(r"^(\d+)\.\s+")
_INNER_FIELD = re.compile(r"([A-Z][^:\[\]]{2,40}):\s*\[")
_ONE_PAGE = re.compile(
    r"\b(?:keep it to one page|in under a page)\b", re.IGNORECASE
)


def _clean_title(text: str) -> str:
    return _NUMBERED_H1.sub("", text).strip()


def _guidance_constraints(guidance: str | None) -> dict | None:
    if guidance and _ONE_PAGE.search(guidance):
        return {"brevity": "terse", "flags": ["one_page"]}
    return None


def parse_default_template(path: Path) -> ParsedWorkbook:
    from engine.structure import DOCX_PARSER_VERSION

    path = Path(path)
    if not path.is_file():
        raise StructureError(f"no template at {path}")
    document = Document(str(path))

    slots: list[dict] = []
    h1_text: str | None = None  # verbatim, for anchors
    h2_text: str | None = None
    h1_num = 0
    section_ordinal = 0
    pending_guidance: str | None = None
    in_front_matter = True
    table_index = -1  # body order among tables, for field locators
    front_meta_done = False

    def anchor() -> str:
        if h1_text and h2_text:
            return f"{h1_text} > {h2_text}"
        return h1_text or "Front matter"

    def human_path() -> str:
        if h1_text and h2_text:
            return f"{_clean_title(h1_text)} > {_clean_title(h2_text)}"
        return _clean_title(h1_text) if h1_text else "Front matter"

    def next_slot_id() -> str:
        nonlocal section_ordinal
        section_ordinal += 1
        base = f"s-h{h1_num:02d}"
        return base if section_ordinal == 1 else f"{base}-{section_ordinal - 1}"

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            para = Paragraph(child, document)
            style = para.style.name if para.style is not None else ""
            text = para.text.strip()
            if not text:
                continue
            if style == "Heading 1":
                m = _NUMBERED_H1.match(text)
                if m:
                    in_front_matter = False
                    h1_text, h2_text = text, None
                    h1_num = int(m.group(1))
                    section_ordinal = 0
                    pending_guidance = None
                    slots.append(_slot(
                        slot_id=f"s-h{h1_num:02d}-hdr",
                        source_mode="firm_default",
                        source_locator={"file": path.name,
                                        "docx_anchor": text},
                        path=human_path(),
                        is_header=True,
                        question_text=_clean_title(text),
                        response_shape="none",
                    ))
                else:
                    in_front_matter = True  # "How to Use This Template"
                continue
            if style == "Heading 2" and not in_front_matter:
                h2_text = text
                slots.append(_slot(
                    slot_id=f"s-h{h1_num:02d}-{field_key(text)}",
                    source_mode="firm_default",
                    source_locator={"file": path.name,
                                    "docx_anchor": anchor()},
                    path=human_path(),
                    parent=f"s-h{h1_num:02d}-hdr",
                    is_header=True,
                    question_text=text,
                    response_shape="none",
                ))
                continue
            if in_front_matter:
                continue
            # Inline bracketed body paragraph ("Payment schedule &
            # terms: [ … ]") → a prose slot in body text.
            if "[" in text and "]" in text:
                slots.append(_slot(
                    slot_id=next_slot_id(),
                    source_mode="firm_default",
                    source_locator={"file": path.name,
                                    "docx_anchor": anchor()},
                    path=human_path(),
                    parent=f"s-h{h1_num:02d}-hdr",
                    question_text=text,
                    response_shape="prose",
                ))
            continue

        if not child.tag.endswith("}tbl"):
            continue
        table = Table(child, document)
        table_index += 1
        first_cell = table.rows[0].cells[0].text.strip() if table.rows else ""

        if in_front_matter:
            # The proposal-metadata table: one addressable record slot.
            if (
                not front_meta_done
                and len(table.columns) == 2
                and len(table.rows) >= 3
                and first_cell.lower() == "field"
            ):
                fields = []
                for r in range(1, len(table.rows)):
                    label = table.rows[r].cells[0].text.strip()
                    if not label:
                        continue
                    fields.append({
                        "key": field_key(label),
                        "label": label,
                        "type": "text",
                        "source_locator": {
                            "table_index": table_index, "row": r, "column": 1,
                        },
                    })
                slots.append(_slot(
                    slot_id="s-front-meta",
                    source_mode="firm_default",
                    source_locator={"file": path.name,
                                    "docx_anchor": "Front matter"},
                    path="Front matter > Proposal metadata",
                    question_text="Proposal metadata",
                    response_shape="record",
                    response_fields=fields,
                ))
                front_meta_done = True
            continue

        if first_cell.startswith("▸"):
            pending_guidance = first_cell
            continue

        if first_cell.startswith("["):
            inner = _INNER_FIELD.findall(first_cell)
            if len(inner) >= 2:
                # The case-study block — repeating structured entry.
                slots.append(_slot(
                    slot_id=next_slot_id(),
                    source_mode="firm_default",
                    source_locator={"file": path.name,
                                    "docx_anchor": anchor()},
                    path=human_path(),
                    parent=f"s-h{h1_num:02d}-hdr",
                    question_text=pending_guidance or first_cell,
                    response_shape="table",
                    response_fields=[
                        {"key": field_key(label), "label": label.strip(),
                         "type": "text",
                         "source_locator": {"table_index": table_index}}
                        for label in inner
                    ],
                ))
            else:
                slots.append(_slot(
                    slot_id=next_slot_id(),
                    source_mode="firm_default",
                    source_locator={"file": path.name,
                                    "docx_anchor": anchor()},
                    path=human_path(),
                    parent=f"s-h{h1_num:02d}-hdr",
                    question_text=pending_guidance,
                    response_shape="prose",
                    constraints=_guidance_constraints(pending_guidance),
                ))
            pending_guidance = None
            continue

        if len(table.columns) >= 2:
            header = [c.text.strip() for c in table.rows[0].cells]
            slots.append(_slot(
                slot_id=next_slot_id(),
                source_mode="firm_default",
                source_locator={"file": path.name,
                                "docx_anchor": anchor()},
                path=human_path(),
                parent=f"s-h{h1_num:02d}-hdr",
                question_text=pending_guidance,
                response_shape="table",
                response_fields=[
                    {"key": field_key(label), "label": label,
                     "type": _field_type(label),
                     "source_locator": {"table_index": table_index,
                                        "column": i}}
                    for i, label in enumerate(header) if label
                ],
            ))
            pending_guidance = None

    if not any(not s.get("is_header") for s in slots):
        raise StructureError(
            f"{path.name}: no answerable slots parsed from the template — "
            "a template with headers only cannot receive a response"
        )

    return ParsedWorkbook(
        file=path.name,
        source_mode="firm_default",
        parser_version=DOCX_PARSER_VERSION,
        source_sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        slots=slots,
    )
